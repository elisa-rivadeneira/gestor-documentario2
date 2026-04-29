"""
API REST para Sistema de Gestión de Correspondencia Institucional
FastAPI + SQLite + Claude IA
"""
import os
import re
import json
import shutil
from dotenv import load_dotenv
load_dotenv(dotenv_path=os.path.join(os.path.dirname(__file__), ".env"))
from datetime import datetime
from typing import Optional, List
from fastapi import FastAPI, Depends, HTTPException, UploadFile, File, Query, Header
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy.orm import Session
from sqlalchemy import or_
import pdfplumber
import io
import openpyxl
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from database import engine, get_db, Base
from models import Documento, Adjunto, Usuario, Contrato, AdjuntoContrato, ComisariaContrato, ExpedienteContrato, PlantillaCarta, CartaGenerada, ConfiguracionSistema, SeguimientoComisaria, SeguimientoCeldaDetalle, RegistroMejora
from schemas import (
    DocumentoCreate, DocumentoUpdate, DocumentoResponse, DocumentoListResponse,
    AdjuntoCreate, AdjuntoResponse, AnalisisIARequest, AnalisisIAResponse,
    LoginRequest, LoginResponse, UsuarioResponse,
    ContratoCreate, ContratoUpdate, ContratoResponse, ContratoListResponse,
    AdjuntoContratoResponse,
    ExpedienteContratoCreate, ExpedienteContratoUpdate, ExpedienteContratoResponse,
    PlantillaCartaCreate, PlantillaCartaResponse,
    GenerarCartaRequest, GenerarCartaResponse, ExportarCartaRequest,
    SeguimientoComisariaResponse, ActualizarCeldaRequest,
    RegistroMejoraCreate, RegistroMejoraUpdate, RegistroMejoraResponse,
    AsistirMejoraRequest, AsistirMejoraResponse
)
from services.ia_service import ia_service, extraer_numero_con_ocr, OCR_DISPONIBLE
from services.auth_service import hash_password, verify_password, create_token, verify_token
from init_users import crear_usuarios_iniciales

# Crear tablas en la base de datos
Base.metadata.create_all(bind=engine)

# Migración: agregar columnas nuevas a contratos si no existen
def migrar_contratos():
    """Agrega columnas nuevas a la tabla contratos si no existen.
    - ruc_contratado: NULL por defecto
    - tipo_contrato: 'mantenimiento' por defecto para registros existentes
    - precio_unitario: NULL por defecto (solo para equipamiento)
    - monto_total: convertir de TEXT a REAL si es necesario
    """
    from sqlalchemy import text
    try:
        with engine.connect() as conn:
            # Verificar columnas existentes
            result = conn.execute(text("PRAGMA table_info(contratos)"))
            columnas_info = result.fetchall()
            columnas = [row[1] for row in columnas_info]

            # Agregar columna ruc_contratado (NULL por defecto)
            if 'ruc_contratado' not in columnas:
                conn.execute(text("ALTER TABLE contratos ADD COLUMN ruc_contratado VARCHAR(11)"))
                conn.commit()
                print("Migración completada: columna ruc_contratado agregada (NULL por defecto)")

            # Agregar columna tipo_contrato
            if 'tipo_contrato' not in columnas:
                conn.execute(text("ALTER TABLE contratos ADD COLUMN tipo_contrato VARCHAR(20)"))
                conn.commit()
                print("Migración completada: columna tipo_contrato agregada")

            # Agregar columna precio_unitario (NULL por defecto, solo para equipamiento)
            if 'precio_unitario' not in columnas:
                conn.execute(text("ALTER TABLE contratos ADD COLUMN precio_unitario REAL"))
                conn.commit()
                print("Migración completada: columna precio_unitario agregada (NULL por defecto)")

            # Agregar columna plazo_dias
            if 'plazo_dias' not in columnas:
                conn.execute(text("ALTER TABLE contratos ADD COLUMN plazo_dias INTEGER"))
                conn.commit()
                print("Migración completada: columna plazo_dias agregada")

            # Agregar columna dias_adicionales (0 por defecto)
            if 'dias_adicionales' not in columnas:
                conn.execute(text("ALTER TABLE contratos ADD COLUMN dias_adicionales INTEGER DEFAULT 0"))
                conn.commit()
                print("Migración completada: columna dias_adicionales agregada (0 por defecto)")

            # Agregar columna estado_ejecucion (PENDIENTE por defecto)
            if 'estado_ejecucion' not in columnas:
                conn.execute(text("ALTER TABLE contratos ADD COLUMN estado_ejecucion VARCHAR(30) DEFAULT 'PENDIENTE'"))
                conn.commit()
                print("Migración completada: columna estado_ejecucion agregada (PENDIENTE por defecto)")

            # Agregar columna tipo_contratado (empresa | consorcio)
            if 'tipo_contratado' not in columnas:
                conn.execute(text("ALTER TABLE contratos ADD COLUMN tipo_contratado VARCHAR(20) DEFAULT 'empresa'"))
                conn.commit()
                print("Migración completada: columna tipo_contratado agregada")

            # Agregar columnas de representante para cartas
            if 'nombre_representante' not in columnas:
                conn.execute(text("ALTER TABLE contratos ADD COLUMN nombre_representante VARCHAR(255)"))
                conn.commit()
                print("Migración completada: columna nombre_representante agregada")

            if 'cargo_representante' not in columnas:
                conn.execute(text("ALTER TABLE contratos ADD COLUMN cargo_representante VARCHAR(255)"))
                conn.commit()
                print("Migración completada: columna cargo_representante agregada")

            if 'email_representante' not in columnas:
                conn.execute(text("ALTER TABLE contratos ADD COLUMN email_representante VARCHAR(255)"))
                conn.commit()
                print("Migración completada: columna email_representante agregada")

            if 'whatsapp_representante' not in columnas:
                conn.execute(text("ALTER TABLE contratos ADD COLUMN whatsapp_representante VARCHAR(20)"))
                conn.commit()
                print("Migración completada: columna whatsapp_representante agregada")

            # Establecer 'mantenimiento' como valor por defecto para contratos existentes sin tipo
            result = conn.execute(text("UPDATE contratos SET tipo_contrato = 'mantenimiento' WHERE tipo_contrato IS NULL"))
            if result.rowcount > 0:
                conn.commit()
                print(f"Migración completada: {result.rowcount} contratos actualizados a tipo 'mantenimiento'")

            # Convertir monto_total de texto a número (extraer solo dígitos y decimales)
            # Solo ejecutar si hay datos con formato texto
            try:
                conn.execute(text("""
                    UPDATE contratos
                    SET monto_total = CAST(
                        REPLACE(REPLACE(REPLACE(monto_total, 'S/', ''), ',', ''), ' ', '')
                        AS REAL
                    )
                    WHERE monto_total IS NOT NULL
                    AND monto_total != ''
                    AND typeof(monto_total) = 'text'
                """))
                conn.commit()
            except:
                pass  # Ignorar si ya es numérico o hay error de conversión

    except Exception as e:
        print(f"Error en migración (puede ignorarse si es nueva instalación): {e}")

migrar_contratos()

def migrar_documentos():
    """
    Agrega columnas nuevas a documentos si no existen.
    - estado: 'enviado' por defecto — todos los registros existentes quedan como enviados.
    """
    from sqlalchemy import text
    try:
        with engine.connect() as conn:
            result = conn.execute(text("PRAGMA table_info(documentos)"))
            columnas = [row[1] for row in result.fetchall()]

            if 'estado' not in columnas:
                conn.execute(text("ALTER TABLE documentos ADD COLUMN estado VARCHAR(20) DEFAULT 'enviado'"))
                conn.commit()
                print("Migración completada: columna 'estado' agregada a documentos (default: 'enviado')")

            if 'archivo_docx' not in columnas:
                conn.execute(text("ALTER TABLE documentos ADD COLUMN archivo_docx VARCHAR(500)"))
                conn.commit()
                print("Migración completada: columna 'archivo_docx' agregada a documentos")

            # Asegurar que todos los registros existentes sin estado queden como 'enviado'
            result = conn.execute(text("UPDATE documentos SET estado='enviado' WHERE estado IS NULL"))
            if result.rowcount > 0:
                conn.commit()
                print(f"Migración completada: {result.rowcount} documentos existentes marcados como 'enviado'")
    except Exception as e:
        print(f"Error en migración documentos: {e}")

migrar_documentos()

def migrar_seguimiento():
    """Agrega columna dossier_monto_merge a seguimiento_comisaria si no existe."""
    from sqlalchemy import text
    try:
        with engine.connect() as conn:
            result = conn.execute(text("PRAGMA table_info(seguimiento_comisaria)"))
            columnas = [row[1] for row in result.fetchall()]
            if 'dossier_monto_merge' not in columnas:
                conn.execute(text("ALTER TABLE seguimiento_comisaria ADD COLUMN dossier_monto_merge INTEGER NOT NULL DEFAULT 0"))
                conn.commit()
                print("Migración completada: columna dossier_monto_merge agregada")
            if 'amp_merge' not in columnas:
                conn.execute(text("ALTER TABLE seguimiento_comisaria ADD COLUMN amp_merge INTEGER NOT NULL DEFAULT 0"))
                conn.commit()
                print("Migración completada: columna amp_merge agregada")
    except Exception as e:
        print(f"Error en migración seguimiento: {e}")

migrar_seguimiento()

# Crear usuarios iniciales si no existen
crear_usuarios_iniciales()

def seed_seguimiento():
    """Pobla la tabla seguimiento_comisaria con los datos del Excel si está vacía."""
    from sqlalchemy.orm import Session as OrmSession
    from database import SessionLocal
    db: OrmSession = SessionLocal()
    try:
        if db.query(SeguimientoComisaria).count() > 0:
            return
        datos = [
            (1,'SAN CAYETANO',1,1,'2026-02-17','2026-02-10','SI','SI','NO','NO','NO','SI','SI','SI','NO','SI','SI','SI','SI',699999.69,'NO','NO',None,None),
            (2,'SAN COSME',1,1,None,'2026-02-18','SI','SI','NO','NO','NO','SI','SI','SI','NO','SI','SI',None,'SI',None,'NO','NO',None,None),
            (3,'COLLIQUE',1,1,None,'2026-03-14','SI','SI','SI','SI','SI','SI','SI','SI','NO','SI','NO','NO','NO',None,'SI','SI',None,None),
            (4,'TAHUANTINSUYO',1,1,None,'2026-03-28','SI','SI','NO','NO','NO','SI','SI','SI','NO','NO','NO','NO','NO',None,'NO','NO',None,None),
            (5,'CHANCAY',1,1,None,'2026-03-28','SI','SI','SI','SI','SI','SI','SI','SI','NO','SI','NO','NO','NO',None,'SI','SI',None,None),
            (6,'JICAMARCA',1,0.8357,None,None,'NO','NO','NO','NO','NO','SI','SI','NO','NO','NO','NO','NO','NO',None,'NO','NO',None,None),
            (7,'PRO',1,1,'2026-03-25',None,'NO','NO','NO','NO','NO','SI','SI','NO','NO','NO','NO','NO','NO',None,'NO','NO',None,None),
            (8,'SAN MARTÍN DE PORRES',1,None,None,None,'NO','NO','NO','NO','NO','SI','SI','NO','NO','NO','NO','NO','NO',None,'NO','NO',None,None),
            (9,'CARABAYLLO',1,1,'2026-03-29','2026-03-24','SI','SI','-','-','-','SI','SI','SI','NO','NO','NO','NO','NO',None,'NO','NO',None,None),
            (10,'LA ENSENADA',1,1,None,'2026-03-29','SI','SI','-','-','-','SI','SI','SI','NO','NO','NO','NO','NO',None,'NO','NO',None,None),
            (11,'MARISCAL CÁCERES',1,1,'2026-04-05','2025-04-05','SI','SI','NO','NO','NO','NO','NO','NO','NO','NO','NO','NO','NO',None,'NO','NO',None,None),
            (12,'SAN ANTONIO DE JICAMARCA',1,1,None,'2025-04-05','SI','SI','NO','NO','NO','NO','NO','NO','NO','NO','NO','NO','NO',None,'NO','NO',None,None),
            (13,'SANTA CLARA',1,1,'2026-04-03','2026-04-03','SI','SI','NO','NO','NO','-','-','-','-','SI','NO','NO','NO',None,'NO','NO',None,None),
            (14,'SANTA ANITA',1,1,None,None,'SI','SI','NO','NO','NO','-','-','-','-','SI','NO','NO','NO',None,'NO','NO',None,'NO QUIERE FIRMAR COMISARIO'),
            (15,'ALFONSO UGARTE',1,None,None,None,'NO','NO','NO','NO','NO','SI','SI','NO','NO','NO','NO','NO','NO',None,'NO','NO',None,None),
            (16,'SAN GENARO',1,1,'2026-03-29',None,'SI','SI','NO','NO','NO','NO','NO','NO','NO','NO','NO','NO','NO',None,'NO','NO',None,None),
            (17,'JOSÉ GÁLVEZ',1,1,None,None,'SI','SI','NO','NO','NO','NO','NO','NO','NO','NO','NO','NO','NO',None,'NO','NO',None,None),
            (18,'VILLA EL SALVADOR',1,0.98,None,None,'NO','NO','NO','NO','NO','NO','NO','NO','NO','NO','NO','NO','NO',None,'NO','NO',None,None),
            (19,'PAMPLONA ALTA II',1,1,None,None,'NO','NO','NO','NO','NO','NO','NO','NO','NO','NO','NO','NO','NO',None,'NO','NO',None,None),
            (20,'CIUDAD Y CAMPO',1,None,None,None,'NO','NO','NO','NO','NO','NO','NO','NO','NO','SI','NO','NO','NO',None,'NO','NO',None,None),
        ]
        for d in datos:
            def parse_date(val):
                if val and isinstance(val, str):
                    from datetime import datetime as dt
                    return dt.strptime(val, '%Y-%m-%d')
                return None
            row = SeguimientoComisaria(
                numero=d[0], comisaria=d[1], avance_programado=d[2], avance_fisico=d[3],
                fecha_fin_contractual=parse_date(d[4]), acta_fecha_firma=parse_date(d[5]),
                acta_revisada=d[6], acta_remitida_ugpe=d[7],
                mod_presentado_ne=d[8], mod_revisado_aprobado=d[9], mod_remitido_ugpe=d[10],
                amp_presentado_ne=d[11], amp_revisado_aprobado=d[12], amp_adenda_firmada=d[13], amp_remitido_ugpe=d[14],
                dossier_presentado_ne=d[15], dossier_revisado_aprobado=d[16], dossier_remitido_ugpe=d[17],
                dossier_remitido_pago=d[18], dossier_monto_pagado=d[19],
                liq_presentado_ne=d[20], liq_revisado_aprobado=d[21], liq_remitido_pago=d[22],
                observaciones=d[23]
            )
            db.add(row)
        db.commit()
        print("Seed seguimiento completado: 20 comisarías cargadas.")
    except Exception as e:
        db.rollback()
        print(f"Error en seed_seguimiento: {e}")
    finally:
        db.close()

seed_seguimiento()

# Crear aplicación FastAPI
app = FastAPI(
    title="Sistema de Gestión de Correspondencia",
    description="MVP para gestión de oficios y cartas institucionales",
    version="1.0.0"
)

# Configurar CORS para permitir frontend local
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # En producción, especificar dominios
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Directorio para archivos subidos (usa variable de entorno en producción)
UPLOAD_DIR = os.getenv("UPLOAD_DIR", os.path.join(os.path.dirname(os.path.dirname(__file__)), "uploads"))
os.makedirs(UPLOAD_DIR, exist_ok=True)

# Montar directorio de uploads
app.mount("/uploads", StaticFiles(directory=UPLOAD_DIR), name="uploads")


# ============================================
# AUTENTICACIÓN
# ============================================

def verificar_admin(authorization: Optional[str] = Header(None)) -> Usuario:
    """
    Dependency para verificar que el usuario está autenticado.
    Extrae el token del header Authorization.
    """
    if not authorization:
        raise HTTPException(status_code=401, detail="No autorizado - Token requerido")

    # Extraer token del header "Bearer <token>"
    try:
        scheme, token = authorization.split()
        if scheme.lower() != "bearer":
            raise HTTPException(status_code=401, detail="Esquema de autorización inválido")
    except ValueError:
        raise HTTPException(status_code=401, detail="Header de autorización inválido")

    # Verificar token
    payload = verify_token(token)
    if not payload:
        raise HTTPException(status_code=401, detail="Token inválido o expirado")

    return payload


def verificar_admin_opcional(authorization: Optional[str] = Header(None)) -> Optional[dict]:
    """
    Dependency opcional - no falla si no hay token, solo retorna None.
    Útil para endpoints que pueden ser públicos o autenticados.
    """
    if not authorization:
        return None

    try:
        scheme, token = authorization.split()
        if scheme.lower() != "bearer":
            return None
        payload = verify_token(token)
        return payload
    except:
        return None


@app.post("/api/login", response_model=LoginResponse)
def login(request: LoginRequest, db: Session = Depends(get_db)):
    """
    Endpoint de login. Verifica credenciales y retorna token JWT.
    """
    # Buscar usuario
    usuario = db.query(Usuario).filter(Usuario.username == request.username).first()

    if not usuario:
        raise HTTPException(status_code=401, detail="Usuario o contraseña incorrectos")

    if not usuario.activo:
        raise HTTPException(status_code=401, detail="Usuario desactivado")

    # Verificar contraseña
    if not verify_password(request.password, usuario.password_hash):
        raise HTTPException(status_code=401, detail="Usuario o contraseña incorrectos")

    # Generar token
    token = create_token(usuario.username, usuario.nombre)

    return LoginResponse(
        token=token,
        usuario=usuario.username,
        nombre=usuario.nombre,
        mensaje="Login exitoso"
    )


@app.get("/api/verificar-token")
def verificar_token_endpoint(admin: dict = Depends(verificar_admin)):
    """
    Verifica si el token actual es válido.
    """
    return {
        "valido": True,
        "usuario": admin.get("sub"),
        "nombre": admin.get("nombre")
    }


# ============================================
# ENDPOINTS DE DOCUMENTOS
# ============================================

@app.get("/api/documentos", response_model=DocumentoListResponse)
def listar_documentos(
    tipo_documento: Optional[str] = Query(None, description="Filtrar por tipo: oficio, carta"),
    direccion: Optional[str] = Query(None, description="Filtrar por dirección: recibido, enviado"),
    busqueda: Optional[str] = Query(None, description="Búsqueda en título, asunto, remitente, destinatario"),
    ordenar_por: Optional[str] = Query(None, description="Ordenar por: numero, fecha"),
    pagina: int = Query(1, ge=1, description="Número de página"),
    por_pagina: int = Query(20, ge=1, le=100, description="Documentos por página"),
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)  # Requiere autenticación
):
    """
    Lista documentos con filtros opcionales y paginación.
    Retorna bandeja unificada de correspondencia.
    """
    query = db.query(Documento)

    # Aplicar filtros
    if tipo_documento:
        query = query.filter(Documento.tipo_documento == tipo_documento)
    if direccion:
        query = query.filter(Documento.direccion == direccion)
    if busqueda:
        busqueda_like = f"%{busqueda}%"
        query = query.filter(
            or_(
                Documento.titulo.ilike(busqueda_like),
                Documento.asunto.ilike(busqueda_like),
                Documento.remitente.ilike(busqueda_like),
                Documento.destinatario.ilike(busqueda_like),
                Documento.numero.ilike(busqueda_like)
            )
        )

    # Contar total
    total = query.count()

    # Aplicar ordenamiento según parámetro
    if ordenar_por == 'fecha':
        # Ordenar por fecha del documento y luego por fecha de subida (más recientes primero)
        documentos = query.order_by(
            Documento.fecha.desc().nullslast(),
            Documento.created_at.desc()
        ).offset((pagina - 1) * por_pagina)\
            .limit(por_pagina)\
            .all()
    else:
        # Ordenar por año y correlativo (más nuevos primero) - default para oficios y cartas nemaec
        documentos = query.order_by(
            Documento.anio_oficio.desc().nullslast(),
            Documento.correlativo_oficio.desc().nullslast(),
            Documento.created_at.desc()
        ).offset((pagina - 1) * por_pagina)\
            .limit(por_pagina)\
            .all()

    return DocumentoListResponse(
        documentos=documentos,
        total=total,
        pagina=pagina,
        por_pagina=por_pagina
    )


@app.post("/api/documentos", response_model=DocumentoResponse, status_code=201)
def crear_documento(
    documento: DocumentoCreate,
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    """
    Crea un nuevo documento (oficio o carta).
    Requiere autenticación de admin.
    """
    # Validar que no exista un documento con el mismo número
    if documento.numero:
        existente = db.query(Documento).filter(Documento.numero == documento.numero).first()
        if existente:
            raise HTTPException(
                status_code=400,
                detail=f"Ya existe un documento con el número {documento.numero}"
            )

    # Validar documento padre si se especifica
    if documento.documento_padre_id:
        padre = db.query(Documento).filter(Documento.id == documento.documento_padre_id).first()
        if not padre:
            raise HTTPException(status_code=404, detail="Documento padre no encontrado")

    # Extraer año y correlativo para ordenamiento
    anio_oficio = None
    correlativo_oficio = None
    if documento.numero:
        import re
        match_correlativo = re.search(r'(\d{3,6})', documento.numero)
        correlativo_oficio = int(match_correlativo.group(1)) if match_correlativo else None
        match_anio = re.search(r'(202[4-9])', documento.numero)
        anio_oficio = int(match_anio.group(1)) if match_anio else None

    db_documento = Documento(
        **documento.model_dump(),
        anio_oficio=anio_oficio,
        correlativo_oficio=correlativo_oficio
    )
    db.add(db_documento)
    db.commit()
    db.refresh(db_documento)
    return db_documento


@app.get("/api/documentos/{documento_id}", response_model=DocumentoResponse)
def obtener_documento(
    documento_id: int,
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)  # Requiere autenticación
):
    """
    Obtiene un documento por su ID con todos sus adjuntos.
    Requiere autenticación.
    """
    documento = db.query(Documento).filter(Documento.id == documento_id).first()
    if not documento:
        raise HTTPException(status_code=404, detail="Documento no encontrado")
    return documento


@app.put("/api/documentos/{documento_id}", response_model=DocumentoResponse)
def actualizar_documento(
    documento_id: int,
    documento_update: DocumentoUpdate,
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    """
    Actualiza un documento existente.
    Requiere autenticación de admin.
    """
    documento = db.query(Documento).filter(Documento.id == documento_id).first()
    if not documento:
        raise HTTPException(status_code=404, detail="Documento no encontrado")

    # Actualizar campos proporcionados
    update_data = documento_update.model_dump(exclude_unset=True)
    for field, value in update_data.items():
        setattr(documento, field, value)

    documento.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(documento)
    return documento


@app.delete("/api/documentos/{documento_id}", status_code=204)
def eliminar_documento(
    documento_id: int,
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    """
    Elimina un documento y sus adjuntos.
    Requiere autenticación de admin.
    """
    documento = db.query(Documento).filter(Documento.id == documento_id).first()
    if not documento:
        raise HTTPException(status_code=404, detail="Documento no encontrado")

    # Eliminar archivo local si existe
    if documento.archivo_local:
        archivo_path = os.path.join(UPLOAD_DIR, documento.archivo_local)
        if os.path.exists(archivo_path):
            os.remove(archivo_path)

    db.delete(documento)
    db.commit()
    return None


@app.get("/api/documentos/{documento_id}/respuestas", response_model=List[DocumentoResponse])
def obtener_respuestas(
    documento_id: int,
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)  # Requiere autenticación
):
    """
    Obtiene las cartas enviadas como respuesta a un documento.
    Requiere autenticación.
    """
    documento = db.query(Documento).filter(Documento.id == documento_id).first()
    if not documento:
        raise HTTPException(status_code=404, detail="Documento no encontrado")

    respuestas = db.query(Documento).filter(
        Documento.documento_padre_id == documento_id
    ).all()

    return respuestas


# ============================================
# ENDPOINTS DE ARCHIVOS
# ============================================

@app.post("/api/documentos/{documento_id}/archivo")
async def subir_archivo(
    documento_id: int,
    archivo: UploadFile = File(...),
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    """
    Sube un archivo PDF a un documento existente.
    Requiere autenticación de admin.
    """
    documento = db.query(Documento).filter(Documento.id == documento_id).first()
    if not documento:
        raise HTTPException(status_code=404, detail="Documento no encontrado")

    # Validar tipo de archivo
    if not archivo.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Solo se permiten archivos PDF")

    # Generar nombre único
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    nombre_archivo = f"{documento_id}_{timestamp}_{archivo.filename}"
    ruta_archivo = os.path.join(UPLOAD_DIR, nombre_archivo)

    # Guardar archivo
    with open(ruta_archivo, "wb") as buffer:
        shutil.copyfileobj(archivo.file, buffer)

    # Actualizar documento
    documento.archivo_local = nombre_archivo
    documento.updated_at = datetime.utcnow()
    db.commit()

    return {
        "mensaje": "Archivo subido exitosamente",
        "archivo": nombre_archivo,
        "ruta": f"/uploads/{nombre_archivo}"
    }


@app.post("/api/subir-temporal")
async def subir_archivo_temporal(
    archivo: UploadFile = File(...),
    admin: dict = Depends(verificar_admin)
):
    """
    Sube un archivo temporalmente para análisis antes de crear el documento.
    Requiere autenticación de admin.
    """
    if not archivo.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Solo se permiten archivos PDF")

    # Generar nombre único temporal
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    nombre_archivo = f"temp_{timestamp}_{archivo.filename}"
    ruta_archivo = os.path.join(UPLOAD_DIR, nombre_archivo)

    # Guardar archivo
    with open(ruta_archivo, "wb") as buffer:
        shutil.copyfileobj(archivo.file, buffer)

    return {
        "mensaje": "Archivo subido temporalmente",
        "archivo": nombre_archivo,
        "ruta": f"/uploads/{nombre_archivo}"
    }


@app.post("/api/documentos/{documento_id}/asociar-archivo")
async def asociar_archivo_temporal(
    documento_id: int,
    nombre_temporal: str = Query(..., description="Nombre del archivo temporal a asociar"),
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    """
    Asocia un archivo temporal ya subido a un documento.
    Renombra el archivo con el ID del documento.
    Requiere autenticación de admin.
    """
    documento = db.query(Documento).filter(Documento.id == documento_id).first()
    if not documento:
        raise HTTPException(status_code=404, detail="Documento no encontrado")

    # Verificar que el archivo temporal existe
    ruta_temporal = os.path.join(UPLOAD_DIR, nombre_temporal)
    if not os.path.exists(ruta_temporal):
        raise HTTPException(status_code=404, detail="Archivo temporal no encontrado")

    # Generar nuevo nombre con el ID del documento
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    # Extraer el nombre original del archivo (quitar prefijo temp_YYYYMMDD_HHMMSS_)
    # El formato temporal es: temp_20260122_203103_NombreOriginal.pdf
    partes = nombre_temporal.split('_', 3)
    nombre_original = partes[3] if len(partes) > 3 else nombre_temporal
    nuevo_nombre = f"{documento_id}_{timestamp}_{nombre_original}"
    ruta_nueva = os.path.join(UPLOAD_DIR, nuevo_nombre)

    # Renombrar archivo
    os.rename(ruta_temporal, ruta_nueva)

    # Actualizar documento
    documento.archivo_local = nuevo_nombre
    documento.updated_at = datetime.utcnow()
    db.commit()

    return {
        "mensaje": "Archivo asociado exitosamente",
        "archivo": nuevo_nombre,
        "ruta": f"/uploads/{nuevo_nombre}"
    }


# ============================================
# ENDPOINTS DE ANÁLISIS IA
# ============================================

@app.post("/api/analizar-ia", response_model=AnalisisIAResponse)
async def analizar_con_ia(
    request: AnalisisIARequest,
    admin: dict = Depends(verificar_admin)
):
    """
    Analiza texto con IA y genera título, asunto y resumen.
    Recibe texto directamente.
    Requiere autenticación de admin.
    """
    if not request.texto:
        raise HTTPException(status_code=400, detail="Se requiere texto para analizar")

    resultado = ia_service.analizar_documento(request.texto)
    return AnalisisIAResponse(**resultado)


@app.post("/api/analizar-archivo/{nombre_archivo}", response_model=AnalisisIAResponse)
async def analizar_archivo_con_ia(
    nombre_archivo: str,
    admin: dict = Depends(verificar_admin)
):
    """
    Extrae texto de un PDF y lo analiza con IA.
    Usa OCR como fallback si no se puede extraer el número de oficio.
    Requiere autenticación de admin.
    """
    ruta_archivo = os.path.join(UPLOAD_DIR, nombre_archivo)

    if not os.path.exists(ruta_archivo):
        raise HTTPException(status_code=404, detail="Archivo no encontrado")

    # Extraer texto del PDF
    try:
        texto = extraer_texto_pdf(ruta_archivo)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error al leer PDF: {str(e)}")

    if not texto or len(texto.strip()) < 50:
        return AnalisisIAResponse(
            numero_oficio="",
            fecha="",
            remitente="",
            destinatario="",
            asunto="",
            resumen="",
            exito=False,
            mensaje="No se pudo extraer suficiente texto del PDF"
        )

    # Agregar el nombre del archivo al inicio del texto para ayudar a la IA
    # El nombre del archivo suele contener el número de oficio
    nombre_original = nombre_archivo.split('_', 3)[-1] if '_' in nombre_archivo else nombre_archivo
    texto_con_nombre = f"NOMBRE DEL ARCHIVO: {nombre_original}\n\n{texto}"

    # Detectar si necesitamos OCR prioritario:
    # 1. Nombre de archivo corto de Windows (contiene ~)
    # 2. El texto tiene "OFICIO N°" pero sin número visible (ej: "OFICIO N° -2026")
    nombre_es_corto_windows = '~' in nombre_original
    texto_sin_numero_encabezado = bool(re.search(r'OFICIO\s*N[°º]?\s*-\s*\d{4}', texto, re.IGNORECASE))

    necesita_ocr_prioritario = nombre_es_corto_windows or texto_sin_numero_encabezado

    # Si necesitamos OCR prioritario y está disponible, extraer número con OCR primero
    numero_ocr = ""
    if necesita_ocr_prioritario and OCR_DISPONIBLE:
        print(f"Nombre corto Windows o texto sin número en encabezado detectado, usando OCR prioritario...")
        numero_ocr = extraer_numero_con_ocr(ruta_archivo)
        print(f"OCR encontró: '{numero_ocr}'")

    # Analizar con IA
    resultado = ia_service.analizar_documento(texto_con_nombre)

    # Si OCR prioritario encontró un número, usarlo (tiene prioridad sobre la IA)
    if numero_ocr:
        resultado["numero_oficio"] = numero_ocr
        resultado["mensaje_whatsapp"] = f"{numero_ocr}\nAsunto: {resultado.get('asunto', '')}\nResumen: {resultado.get('resumen', '')}"
        resultado["mensaje"] = "Análisis completado (número extraído con OCR)"
    else:
        # Verificar si el número de oficio tiene el formato correcto
        # Acepta 5-6 dígitos para oficios O 1-6 dígitos para cartas NEMAEC O carta genérica
        numero_actual = resultado.get("numero_oficio", "")
        es_nemaec = "NEMAEC" in numero_actual.upper()
        es_carta = "CARTA" in numero_actual.upper()
        tiene_numero_valido = bool(re.search(r'\d{5,6}', numero_actual)) or (es_nemaec and bool(re.search(r'\d{1,6}', numero_actual))) or (es_carta and bool(re.search(r'\d{1,6}', numero_actual)))

        # Si no se encontró número válido y OCR está disponible, intentar con OCR
        if not tiene_numero_valido and OCR_DISPONIBLE:
            print(f"Número de oficio incompleto o no encontrado: '{numero_actual}', intentando OCR...")
            numero_ocr = extraer_numero_con_ocr(ruta_archivo)
            if numero_ocr:
                resultado["numero_oficio"] = numero_ocr
                # Actualizar mensaje WhatsApp con el número encontrado por OCR
                resultado["mensaje_whatsapp"] = f"{numero_ocr}\nAsunto: {resultado.get('asunto', '')}\nResumen: {resultado.get('resumen', '')}"
                resultado["mensaje"] = "Análisis completado (número extraído con OCR)"

    return AnalisisIAResponse(**resultado)


def extraer_texto_pdf(ruta: str) -> str:
    """
    Extrae texto de un archivo PDF usando pdfplumber (mejor extracción).
    """
    texto = ""
    with pdfplumber.open(ruta) as pdf:
        for page in pdf.pages:
            texto += page.extract_text() or ""
    return texto


# ============================================
# ENDPOINTS DE ADJUNTOS
# ============================================

@app.post("/api/documentos/{documento_id}/adjuntos", response_model=AdjuntoResponse)
async def agregar_adjunto(
    documento_id: int,
    archivo: UploadFile = File(None),
    enlace_drive: str = Query(None),
    nombre: str = Query(None),
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    """
    Agrega un adjunto a un documento (archivo o enlace Drive).
    Requiere autenticación de admin.
    """
    documento = db.query(Documento).filter(Documento.id == documento_id).first()
    if not documento:
        raise HTTPException(status_code=404, detail="Documento no encontrado")

    adjunto_data = {"documento_id": documento_id}

    if archivo:
        # Subir archivo
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        nombre_archivo = f"adj_{documento_id}_{timestamp}_{archivo.filename}"
        ruta_archivo = os.path.join(UPLOAD_DIR, nombre_archivo)

        with open(ruta_archivo, "wb") as buffer:
            shutil.copyfileobj(archivo.file, buffer)

        adjunto_data["archivo_local"] = nombre_archivo
        adjunto_data["nombre"] = nombre or archivo.filename
    elif enlace_drive:
        adjunto_data["enlace_drive"] = enlace_drive
        adjunto_data["nombre"] = nombre or "Enlace Drive"
    else:
        raise HTTPException(status_code=400, detail="Se requiere archivo o enlace Drive")

    adjunto = Adjunto(**adjunto_data)
    db.add(adjunto)
    db.commit()
    db.refresh(adjunto)
    return adjunto


@app.delete("/api/adjuntos/{adjunto_id}", status_code=204)
def eliminar_adjunto(
    adjunto_id: int,
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    """
    Elimina un adjunto.
    Requiere autenticación de admin.
    """
    adjunto = db.query(Adjunto).filter(Adjunto.id == adjunto_id).first()
    if not adjunto:
        raise HTTPException(status_code=404, detail="Adjunto no encontrado")

    # Eliminar archivo si existe
    if adjunto.archivo_local:
        ruta = os.path.join(UPLOAD_DIR, adjunto.archivo_local)
        if os.path.exists(ruta):
            os.remove(ruta)

    db.delete(adjunto)
    db.commit()
    return None


# ============================================
# ENDPOINTS DE CONTRATOS
# ============================================

@app.get("/api/contratos", response_model=ContratoListResponse)
def listar_contratos(
    busqueda: Optional[str] = Query(None, description="Búsqueda en contratante, contratado, item, asunto, número"),
    tipo_contrato: Optional[str] = Query(None, description="Filtrar por tipo: equipamiento,mantenimiento (separados por coma)"),
    pagina: int = Query(1, ge=1),
    por_pagina: int = Query(20, ge=1, le=100),
    db: Session = Depends(get_db)
):
    """Lista contratos con búsqueda opcional y paginación."""
    query = db.query(Contrato)

    # Filtrar por tipo de contrato
    if tipo_contrato:
        tipos = [t.strip() for t in tipo_contrato.split(',') if t.strip()]
        if tipos:
            query = query.filter(Contrato.tipo_contrato.in_(tipos))

    if busqueda:
        busqueda_like = f"%{busqueda}%"
        query = query.filter(
            or_(
                Contrato.contratante.ilike(busqueda_like),
                Contrato.contratado.ilike(busqueda_like),
                Contrato.item_contratado.ilike(busqueda_like),
                Contrato.asunto.ilike(busqueda_like),
                Contrato.numero.ilike(busqueda_like)
            )
        )

    total = query.count()
    contratos = query.order_by(Contrato.created_at.desc())\
        .offset((pagina - 1) * por_pagina)\
        .limit(por_pagina)\
        .all()

    return ContratoListResponse(
        contratos=contratos,
        total=total,
        pagina=pagina,
        por_pagina=por_pagina
    )


@app.post("/api/contratos", response_model=ContratoResponse, status_code=201)
def crear_contrato(
    contrato: ContratoCreate,
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    """Crea un nuevo contrato. Requiere autenticación."""
    # Extraer comisarías del payload
    comisarias_data = contrato.comisarias
    contrato_dict = contrato.model_dump(exclude={'comisarias'})

    # Crear el contrato
    db_contrato = Contrato(**contrato_dict)
    db.add(db_contrato)
    db.commit()
    db.refresh(db_contrato)

    # Si es mantenimiento y hay comisarías, crearlas
    if contrato.tipo_contrato == 'mantenimiento' and comisarias_data:
        for comisaria in comisarias_data:
            db_comisaria = ComisariaContrato(
                contrato_id=db_contrato.id,
                nombre_cpnp=comisaria.nombre_cpnp,
                monto=comisaria.monto
            )
            db.add(db_comisaria)
        db.commit()
        db.refresh(db_contrato)

    return db_contrato


@app.get("/api/contratos/{contrato_id}", response_model=ContratoResponse)
def obtener_contrato(contrato_id: int, db: Session = Depends(get_db)):
    """Obtiene un contrato por su ID con todos sus adjuntos."""
    contrato = db.query(Contrato).filter(Contrato.id == contrato_id).first()
    if not contrato:
        raise HTTPException(status_code=404, detail="Contrato no encontrado")
    return contrato


@app.put("/api/contratos/{contrato_id}", response_model=ContratoResponse)
def actualizar_contrato(
    contrato_id: int,
    contrato_update: ContratoUpdate,
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    """Actualiza un contrato existente. Requiere autenticación."""
    contrato = db.query(Contrato).filter(Contrato.id == contrato_id).first()
    if not contrato:
        raise HTTPException(status_code=404, detail="Contrato no encontrado")

    # Extraer comisarías del payload
    comisarias_data = contrato_update.comisarias
    update_data = contrato_update.model_dump(exclude_unset=True, exclude={'comisarias'})

    for field, value in update_data.items():
        setattr(contrato, field, value)

    # Si se envían comisarías, reemplazar las existentes
    if comisarias_data is not None:
        # Eliminar comisarías existentes
        db.query(ComisariaContrato).filter(ComisariaContrato.contrato_id == contrato_id).delete()

        # Crear nuevas comisarías
        for comisaria in comisarias_data:
            db_comisaria = ComisariaContrato(
                contrato_id=contrato_id,
                nombre_cpnp=comisaria.nombre_cpnp,
                monto=comisaria.monto
            )
            db.add(db_comisaria)

    contrato.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(contrato)
    return contrato


@app.delete("/api/contratos/{contrato_id}", status_code=204)
def eliminar_contrato(
    contrato_id: int,
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    """Elimina un contrato y sus adjuntos. Requiere autenticación."""
    contrato = db.query(Contrato).filter(Contrato.id == contrato_id).first()
    if not contrato:
        raise HTTPException(status_code=404, detail="Contrato no encontrado")

    if contrato.archivo_local:
        archivo_path = os.path.join(UPLOAD_DIR, contrato.archivo_local)
        if os.path.exists(archivo_path):
            os.remove(archivo_path)

    # Eliminar archivos de adjuntos
    for adj in contrato.adjuntos:
        if adj.archivo_local:
            ruta = os.path.join(UPLOAD_DIR, adj.archivo_local)
            if os.path.exists(ruta):
                os.remove(ruta)

    db.delete(contrato)
    db.commit()
    return None


@app.post("/api/contratos/{contrato_id}/asociar-archivo")
async def asociar_archivo_contrato(
    contrato_id: int,
    nombre_temporal: str = Query(..., description="Nombre del archivo temporal a asociar"),
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    """Asocia un archivo temporal ya subido a un contrato."""
    contrato = db.query(Contrato).filter(Contrato.id == contrato_id).first()
    if not contrato:
        raise HTTPException(status_code=404, detail="Contrato no encontrado")

    ruta_temporal = os.path.join(UPLOAD_DIR, nombre_temporal)
    if not os.path.exists(ruta_temporal):
        raise HTTPException(status_code=404, detail="Archivo temporal no encontrado")

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    partes = nombre_temporal.split('_', 3)
    nombre_original = partes[3] if len(partes) > 3 else nombre_temporal
    nuevo_nombre = f"contrato_{contrato_id}_{timestamp}_{nombre_original}"
    ruta_nueva = os.path.join(UPLOAD_DIR, nuevo_nombre)

    # Eliminar archivo anterior si existe
    if contrato.archivo_local:
        ruta_anterior = os.path.join(UPLOAD_DIR, contrato.archivo_local)
        if os.path.exists(ruta_anterior):
            os.remove(ruta_anterior)

    os.rename(ruta_temporal, ruta_nueva)

    contrato.archivo_local = nuevo_nombre
    contrato.updated_at = datetime.utcnow()
    db.commit()

    return {
        "mensaje": "Archivo asociado exitosamente",
        "archivo": nuevo_nombre,
        "ruta": f"/uploads/{nuevo_nombre}"
    }


@app.post("/api/contratos/{contrato_id}/adjuntos", response_model=AdjuntoContratoResponse)
async def agregar_adjunto_contrato(
    contrato_id: int,
    archivo: UploadFile = File(None),
    enlace_drive: str = Query(None),
    nombre: str = Query(None),
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    """Agrega un adjunto a un contrato (archivo o enlace Drive)."""
    contrato = db.query(Contrato).filter(Contrato.id == contrato_id).first()
    if not contrato:
        raise HTTPException(status_code=404, detail="Contrato no encontrado")

    adjunto_data = {"contrato_id": contrato_id}

    if archivo:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        nombre_archivo = f"adj_contrato_{contrato_id}_{timestamp}_{archivo.filename}"
        ruta_archivo = os.path.join(UPLOAD_DIR, nombre_archivo)

        with open(ruta_archivo, "wb") as buffer:
            shutil.copyfileobj(archivo.file, buffer)

        adjunto_data["archivo_local"] = nombre_archivo
        adjunto_data["nombre"] = nombre or archivo.filename
    elif enlace_drive:
        adjunto_data["enlace_drive"] = enlace_drive
        adjunto_data["nombre"] = nombre or "Enlace Drive"
    else:
        raise HTTPException(status_code=400, detail="Se requiere archivo o enlace Drive")

    adjunto = AdjuntoContrato(**adjunto_data)
    db.add(adjunto)
    db.commit()
    db.refresh(adjunto)
    return adjunto


@app.delete("/api/adjuntos-contrato/{adjunto_id}", status_code=204)
def eliminar_adjunto_contrato(
    adjunto_id: int,
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    """Elimina un adjunto de contrato."""
    adjunto = db.query(AdjuntoContrato).filter(AdjuntoContrato.id == adjunto_id).first()
    if not adjunto:
        raise HTTPException(status_code=404, detail="Adjunto no encontrado")

    if adjunto.archivo_local:
        ruta = os.path.join(UPLOAD_DIR, adjunto.archivo_local)
        if os.path.exists(ruta):
            os.remove(ruta)

    db.delete(adjunto)
    db.commit()
    return None


# ============================================
# ENDPOINT DE RESTAURACIÓN DE UPLOADS
# ============================================

@app.post("/api/restaurar-uploads")
async def restaurar_uploads(
    archivo: UploadFile = File(...),
    admin: dict = Depends(verificar_admin)
):
    """
    Restaura archivos de uploads desde un ZIP.
    Solo admin puede usar este endpoint.
    """
    import zipfile
    import io

    if not archivo.filename.lower().endswith('.zip'):
        raise HTTPException(status_code=400, detail="Solo se permiten archivos ZIP")

    # Leer el ZIP
    contenido = await archivo.read()

    try:
        archivos_restaurados = []
        with zipfile.ZipFile(io.BytesIO(contenido), 'r') as zip_ref:
            for nombre in zip_ref.namelist():
                # Ignorar directorios y archivos ocultos
                if nombre.endswith('/') or nombre.startswith('.') or '/' in nombre:
                    continue

                # Extraer solo archivos PDF
                if nombre.lower().endswith('.pdf'):
                    ruta_destino = os.path.join(UPLOAD_DIR, nombre)
                    # Solo extraer si no existe (no sobrescribir)
                    if not os.path.exists(ruta_destino):
                        with zip_ref.open(nombre) as src, open(ruta_destino, 'wb') as dst:
                            dst.write(src.read())
                        archivos_restaurados.append(nombre)

        return {
            "mensaje": f"Restauración completada",
            "archivos_restaurados": len(archivos_restaurados),
            "lista": archivos_restaurados[:20]  # Mostrar máximo 20
        }
    except zipfile.BadZipFile:
        raise HTTPException(status_code=400, detail="Archivo ZIP inválido")


# ============================================
# ENDPOINT DE SALUD
# ============================================

@app.get("/api/health")
def health_check():
    """Verificar estado del servidor"""
    return {"status": "ok", "timestamp": datetime.utcnow().isoformat()}


# ============================================
# ENDPOINT EXTRAER REFERENCIA DESDE PDF
# ============================================

@app.post("/api/extraer-referencia-pdf")
async def extraer_referencia_pdf(
    archivo: UploadFile = File(None),
    documento_id: int = Query(None),
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    """
    Extrae información de referencia (tipo, número, fecha, asunto) desde la
    primera página de un PDF, usando IA.
    Acepta un archivo subido O un documento_id existente en el gestor.
    """
    from openai import OpenAI
    import tempfile

    ruta_temp = None
    ruta_pdf = None

    try:
        # Obtener ruta al PDF
        if archivo:
            suffix = os.path.splitext(archivo.filename or ".pdf")[1] or ".pdf"
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix, dir=UPLOAD_DIR)
            contenido = await archivo.read()
            tmp.write(contenido)
            tmp.close()
            ruta_temp = tmp.name
            ruta_pdf = ruta_temp
        elif documento_id:
            doc = db.query(Documento).filter(Documento.id == documento_id).first()
            if not doc:
                raise HTTPException(status_code=404, detail="Documento no encontrado")
            if not doc.archivo_local:
                raise HTTPException(status_code=400, detail="El documento no tiene archivo asociado")
            ruta_pdf = os.path.join(UPLOAD_DIR, doc.archivo_local)
            if not os.path.exists(ruta_pdf):
                raise HTTPException(status_code=404, detail="Archivo PDF no encontrado en disco")
        else:
            raise HTTPException(status_code=400, detail="Se requiere 'archivo' o 'documento_id'")

        # Extraer solo primera página
        texto_primera_pagina = ""
        try:
            with pdfplumber.open(ruta_pdf) as pdf:
                if pdf.pages:
                    texto_primera_pagina = pdf.pages[0].extract_text() or ""
        except Exception as e:
            raise HTTPException(status_code=422, detail=f"No se pudo leer el PDF: {e}")

        if not texto_primera_pagina.strip():
            return {"exito": False, "mensaje": "No se pudo extraer texto del PDF (puede ser imagen escaneada)"}

        # Llamar a IA para extraer campos
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise HTTPException(status_code=503, detail="API de IA no configurada")

        client = OpenAI(api_key=api_key)
        prompt = f"""Analiza el siguiente texto de la primera página de un documento y extrae:
- tipo_doc: el tipo de documento (Carta, Oficio, Convenio, Informe, Acta, Resolución, Contrato u Otro)
- numero: el número o código del documento (solo el número, sin "N°" ni palabras extra)
- fecha: la fecha del documento en formato YYYY-MM-DD (si no hay fecha clara, dejar vacío)
- asunto: el asunto o tema principal del documento en una línea corta (máximo 100 caracteres)

Responde ÚNICAMENTE con JSON sin markdown:
{{"tipo_doc": "...", "numero": "...", "fecha": "...", "asunto": "..."}}

TEXTO DEL DOCUMENTO:
{texto_primera_pagina[:2000]}"""

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=200,
        )
        raw = response.choices[0].message.content.strip()
        if raw.startswith("```"):
            raw = re.sub(r'^```[a-z]*\n?', '', raw)
            raw = re.sub(r'\n?```$', '', raw)
        datos = json.loads(raw)

        return {
            "exito": True,
            "tipo_doc": datos.get("tipo_doc", "Documento"),
            "numero": datos.get("numero", ""),
            "fecha": datos.get("fecha", ""),
            "asunto": datos.get("asunto", ""),
        }

    finally:
        if ruta_temp and os.path.exists(ruta_temp):
            os.remove(ruta_temp)


# ============================================
# ENDPOINT DE CONSULTA RUC (SUNAT)
# ============================================

@app.get("/api/consultar-ruc/{ruc}")
async def consultar_ruc(ruc: str, admin: dict = Depends(verificar_admin)):
    """
    Consulta la razón social de un RUC en SUNAT.
    Usa la API gratuita de apis.net.pe
    Requiere autenticación.
    """
    import httpx

    # Validar formato de RUC (11 dígitos)
    if not ruc.isdigit() or len(ruc) != 11:
        raise HTTPException(status_code=400, detail="El RUC debe tener exactamente 11 dígitos")

    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            # Usar API gratuita de apis.net.pe
            response = await client.get(f"https://api.apis.net.pe/v1/ruc?numero={ruc}")

            if response.status_code == 200:
                data = response.json()
                return {
                    "exito": True,
                    "ruc": data.get("numeroDocumento", ruc),
                    "razon_social": data.get("nombre", ""),
                    "estado": data.get("estado", ""),
                    "condicion": data.get("condicion", ""),
                    "direccion": data.get("direccion", ""),
                    "departamento": data.get("departamento", ""),
                    "provincia": data.get("provincia", ""),
                    "distrito": data.get("distrito", "")
                }
            elif response.status_code == 404:
                return {
                    "exito": False,
                    "mensaje": "RUC no encontrado en SUNAT"
                }
            else:
                return {
                    "exito": False,
                    "mensaje": f"Error al consultar SUNAT: {response.status_code}"
                }
    except httpx.TimeoutException:
        return {
            "exito": False,
            "mensaje": "Tiempo de espera agotado al consultar SUNAT"
        }
    except Exception as e:
        return {
            "exito": False,
            "mensaje": f"Error de conexión: {str(e)}"
        }


# ============================================
# EXPEDIENTE POR CONTRATO
# ============================================

def _mover_archivo_expediente(nombre_temporal: str, contrato_id: int) -> str:
    """Mueve un archivo temporal a su nombre definitivo para el expediente."""
    ruta_temporal = os.path.join(UPLOAD_DIR, nombre_temporal)
    if not os.path.exists(ruta_temporal):
        raise HTTPException(status_code=404, detail="Archivo temporal no encontrado")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    partes = nombre_temporal.split('_', 3)
    nombre_original = partes[3] if len(partes) > 3 else nombre_temporal
    nuevo_nombre = f"exp_{contrato_id}_{timestamp}_{nombre_original}"
    os.rename(ruta_temporal, os.path.join(UPLOAD_DIR, nuevo_nombre))
    return nuevo_nombre


@app.get("/api/contratos/{contrato_id}/expediente", response_model=List[ExpedienteContratoResponse])
def listar_expediente_contrato(
    contrato_id: int,
    db: Session = Depends(get_db)
):
    """Lista todo el expediente de un contrato, ordenado por fecha asc (cronológico)."""
    contrato = db.query(Contrato).filter(Contrato.id == contrato_id).first()
    if not contrato:
        raise HTTPException(status_code=404, detail="Contrato no encontrado")
    items = (
        db.query(ExpedienteContrato)
        .filter(ExpedienteContrato.contrato_id == contrato_id)
        .order_by(ExpedienteContrato.fecha.asc().nullslast(), ExpedienteContrato.created_at.asc())
        .all()
    )
    return items


@app.post("/api/contratos/{contrato_id}/expediente", response_model=ExpedienteContratoResponse)
def crear_expediente_contrato(
    contrato_id: int,
    data: ExpedienteContratoCreate,
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    """Agrega un documento al expediente de un contrato."""
    contrato = db.query(Contrato).filter(Contrato.id == contrato_id).first()
    if not contrato:
        raise HTTPException(status_code=404, detail="Contrato no encontrado")

    archivo_local = None
    if data.archivo_temporal:
        archivo_local = _mover_archivo_expediente(data.archivo_temporal, contrato_id)

    item = ExpedienteContrato(
        contrato_id=contrato_id,
        tipo_doc=data.tipo_doc,
        numero=data.numero,
        fecha=data.fecha,
        asunto=data.asunto,
        archivo_local=archivo_local,
        enlace_drive=data.enlace_drive,
        notas=data.notas,
    )
    db.add(item)
    db.commit()
    db.refresh(item)
    return item


@app.put("/api/expediente/{item_id}", response_model=ExpedienteContratoResponse)
def actualizar_expediente(
    item_id: int,
    data: ExpedienteContratoUpdate,
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    """Actualiza un documento del expediente."""
    item = db.query(ExpedienteContrato).filter(ExpedienteContrato.id == item_id).first()
    if not item:
        raise HTTPException(status_code=404, detail="Documento no encontrado")

    if data.archivo_temporal:
        if item.archivo_local:
            ruta_anterior = os.path.join(UPLOAD_DIR, item.archivo_local)
            if os.path.exists(ruta_anterior):
                os.remove(ruta_anterior)
        item.archivo_local = _mover_archivo_expediente(data.archivo_temporal, item.contrato_id)

    update_data = data.model_dump(exclude_unset=True, exclude={'archivo_temporal'})
    for field, value in update_data.items():
        setattr(item, field, value)

    db.commit()
    db.refresh(item)
    return item


@app.delete("/api/expediente/{item_id}", status_code=204)
def eliminar_expediente(
    item_id: int,
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    """Elimina un documento del expediente."""
    item = db.query(ExpedienteContrato).filter(ExpedienteContrato.id == item_id).first()
    if not item:
        raise HTTPException(status_code=404, detail="Documento no encontrado")
    if item.archivo_local:
        ruta = os.path.join(UPLOAD_DIR, item.archivo_local)
        if os.path.exists(ruta):
            os.remove(ruta)
    db.delete(item)
    db.commit()
    return None


# ============================================
# PLANTILLAS DE CARTA
# ============================================

@app.get("/api/plantillas-carta", response_model=List[PlantillaCartaResponse])
def listar_plantillas(
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    """Lista todas las plantillas de carta."""
    return db.query(PlantillaCarta).order_by(PlantillaCarta.created_at.desc()).all()


@app.post("/api/plantillas-carta", response_model=PlantillaCartaResponse)
def crear_plantilla(
    nombre: str = Query(...),
    descripcion: Optional[str] = Query(None),
    archivo: UploadFile = File(...),
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    """Crea una nueva plantilla de carta subiendo un archivo .docx de referencia."""
    # Guardar el archivo
    ext = os.path.splitext(archivo.filename)[1].lower()
    if ext not in ['.docx', '.doc']:
        raise HTTPException(status_code=400, detail="Solo se aceptan archivos .docx o .doc")

    carpeta = os.path.join(UPLOAD_DIR, "plantillas")
    os.makedirs(carpeta, exist_ok=True)
    nombre_archivo = f"plantilla_{datetime.now().strftime('%Y%m%d_%H%M%S')}{ext}"
    ruta = os.path.join(carpeta, nombre_archivo)

    with open(ruta, "wb") as f:
        f.write(archivo.file.read())

    plantilla = PlantillaCarta(
        nombre=nombre,
        descripcion=descripcion,
        archivo_local=f"plantillas/{nombre_archivo}",
    )
    db.add(plantilla)
    db.commit()
    db.refresh(plantilla)
    return plantilla


@app.delete("/api/plantillas-carta/{plantilla_id}", status_code=204)
def eliminar_plantilla(
    plantilla_id: int,
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    """Elimina una plantilla de carta."""
    plantilla = db.query(PlantillaCarta).filter(PlantillaCarta.id == plantilla_id).first()
    if not plantilla:
        raise HTTPException(status_code=404, detail="Plantilla no encontrada")
    if plantilla.archivo_local:
        ruta = os.path.join(UPLOAD_DIR, plantilla.archivo_local)
        if os.path.exists(ruta):
            os.remove(ruta)
    db.delete(plantilla)
    db.commit()
    return None


# ============================================
# MEMBRETE
# ============================================

@app.get("/api/membrete")
def get_membrete(db: Session = Depends(get_db)):
    """Retorna la ruta del membrete actual si existe."""
    cfg = db.query(ConfiguracionSistema).filter(ConfiguracionSistema.clave == "membrete_archivo").first()
    if cfg and cfg.valor:
        return {"archivo": cfg.valor, "url": f"/uploads/{cfg.valor}"}
    return {"archivo": None, "url": None}


@app.post("/api/membrete")
async def subir_membrete(
    archivo: UploadFile = File(...),
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    """Sube o reemplaza el membrete (.docx). Solo puede haber uno."""
    if not archivo.filename.lower().endswith('.docx'):
        raise HTTPException(status_code=400, detail="Solo se aceptan archivos .docx")

    # Eliminar membrete anterior si existe
    cfg = db.query(ConfiguracionSistema).filter(ConfiguracionSistema.clave == "membrete_archivo").first()
    if cfg and cfg.valor:
        ruta_anterior = os.path.join(UPLOAD_DIR, cfg.valor)
        if os.path.exists(ruta_anterior):
            os.remove(ruta_anterior)

    nombre_archivo = f"membrete_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    ruta = os.path.join(UPLOAD_DIR, nombre_archivo)
    contenido = await archivo.read()
    with open(ruta, "wb") as f:
        f.write(contenido)

    if cfg:
        cfg.valor = nombre_archivo
    else:
        cfg = ConfiguracionSistema(clave="membrete_archivo", valor=nombre_archivo)
        db.add(cfg)
    db.commit()
    return {"archivo": nombre_archivo, "url": f"/uploads/{nombre_archivo}"}


@app.delete("/api/membrete", status_code=204)
def eliminar_membrete(
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    """Elimina el membrete actual."""
    cfg = db.query(ConfiguracionSistema).filter(ConfiguracionSistema.clave == "membrete_archivo").first()
    if cfg and cfg.valor:
        ruta = os.path.join(UPLOAD_DIR, cfg.valor)
        if os.path.exists(ruta):
            os.remove(ruta)
        cfg.valor = None
        db.commit()
    return None


@app.get("/api/configuracion/firma")
def get_configuracion_firma(db: Session = Depends(get_db)):
    def get_cfg(clave, default=""):
        row = db.query(ConfiguracionSistema).filter(ConfiguracionSistema.clave == clave).first()
        return row.valor if row and row.valor else default
    archivo = get_cfg("firma_imagen")
    return {
        "nombre": get_cfg("firma_nombre"),
        "cargo": get_cfg("firma_cargo"),
        "imagen_archivo": archivo,
        "imagen_url": f"/uploads/{archivo}" if archivo else None,
    }


@app.post("/api/configuracion/firma")
def guardar_configuracion_firma(
    data: dict,
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    def set_cfg(clave, valor):
        row = db.query(ConfiguracionSistema).filter(ConfiguracionSistema.clave == clave).first()
        if row:
            row.valor = str(valor)
        else:
            db.add(ConfiguracionSistema(clave=clave, valor=str(valor)))

    set_cfg("firma_nombre", data.get("nombre", ""))
    set_cfg("firma_cargo", data.get("cargo", ""))
    db.commit()
    return {"ok": True}


@app.post("/api/configuracion/firma/imagen")
async def subir_firma_imagen(
    archivo: UploadFile = File(...),
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    if not archivo.filename.lower().endswith(('.png', '.jpg', '.jpeg')):
        raise HTTPException(status_code=400, detail="Solo se aceptan PNG o JPG")

    cfg = db.query(ConfiguracionSistema).filter(ConfiguracionSistema.clave == "firma_imagen").first()
    if cfg and cfg.valor:
        ruta_anterior = os.path.join(UPLOAD_DIR, cfg.valor)
        if os.path.exists(ruta_anterior):
            os.remove(ruta_anterior)

    ext = os.path.splitext(archivo.filename)[1].lower()
    nombre_archivo = f"firma_{datetime.now().strftime('%Y%m%d%H%M%S')}{ext}"
    ruta = os.path.join(UPLOAD_DIR, nombre_archivo)
    contenido = await archivo.read()
    with open(ruta, "wb") as f:
        f.write(contenido)

    if cfg:
        cfg.valor = nombre_archivo
    else:
        db.add(ConfiguracionSistema(clave="firma_imagen", valor=nombre_archivo))
    db.commit()
    return {"archivo": nombre_archivo, "url": f"/uploads/{nombre_archivo}"}


@app.delete("/api/configuracion/firma/imagen", status_code=204)
def eliminar_firma_imagen(
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    cfg = db.query(ConfiguracionSistema).filter(ConfiguracionSistema.clave == "firma_imagen").first()
    if cfg and cfg.valor:
        ruta = os.path.join(UPLOAD_DIR, cfg.valor)
        if os.path.exists(ruta):
            os.remove(ruta)
        cfg.valor = None
        db.commit()
    return None


@app.get("/api/configuracion/numeracion")
def get_configuracion_numeracion(db: Session = Depends(get_db)):
    """Retorna la configuración de numeración de cartas."""
    def get_cfg(clave, default):
        row = db.query(ConfiguracionSistema).filter(ConfiguracionSistema.clave == clave).first()
        return row.valor if row and row.valor else default

    return {
        "sufijo": get_cfg("carta_sufijo", ""),
        "digitos": int(get_cfg("carta_digitos", "6")),
    }


@app.post("/api/configuracion/numeracion")
def guardar_configuracion_numeracion(
    data: dict,
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    """Guarda la configuración de numeración de cartas."""
    def set_cfg(clave, valor):
        row = db.query(ConfiguracionSistema).filter(ConfiguracionSistema.clave == clave).first()
        if row:
            row.valor = str(valor)
        else:
            db.add(ConfiguracionSistema(clave=clave, valor=str(valor)))

    set_cfg("carta_sufijo", data.get("sufijo", ""))
    set_cfg("carta_digitos", str(data.get("digitos", 6)))
    db.commit()
    return {"ok": True}


# ============================================
# GENERADOR DE CARTAS CON IA
# ============================================

def _obtener_siguiente_numero_carta(db: Session) -> tuple:
    """
    Obtiene el siguiente número correlativo para una carta.
    Lee configuración de sufijo y dígitos desde configuracion_sistema.
    Retorna (numero_correlativo: int, numero_completo: str).
    """
    anio = datetime.now().year

    # Leer configuración
    def get_cfg(clave, default):
        row = db.query(ConfiguracionSistema).filter(ConfiguracionSistema.clave == clave).first()
        return row.valor if row and row.valor else default

    sufijo = get_cfg("carta_sufijo", "").strip()
    digitos = int(get_cfg("carta_digitos", "6"))

    ultima = (
        db.query(CartaGenerada)
        .filter(CartaGenerada.anio == anio)
        .order_by(CartaGenerada.numero_correlativo.desc())
        .first()
    )
    siguiente = (ultima.numero_correlativo + 1) if ultima else 1
    correlativo_str = str(siguiente).zfill(digitos)

    if sufijo:
        numero_completo = f"Carta N° {correlativo_str}-{anio}-{sufijo}"
    else:
        numero_completo = f"Carta N° {correlativo_str}-{anio}"

    return siguiente, numero_completo


def _leer_plantilla_docx(ruta: str) -> str:
    """Extrae texto de un .docx para usarlo como referencia de estructura."""
    try:
        import zipfile
        import xml.etree.ElementTree as ET
        with zipfile.ZipFile(ruta) as z:
            with z.open('word/document.xml') as f:
                tree = ET.parse(f)
                root = tree.getroot()
                texts = []
                for para in root.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
                    runs = para.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                    line = ''.join(r.text or '' for r in runs)
                    if line.strip():
                        texts.append(line)
        return '\n'.join(texts[:60])  # Primeras 60 líneas como referencia
    except Exception as e:
        print(f"Error leyendo plantilla: {e}")
        return ""


@app.post("/api/generar-carta", response_model=GenerarCartaResponse)
def generar_carta_ia(
    request: GenerarCartaRequest,
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    """
    Genera una propuesta de carta usando IA.
    Toma datos del contrato como destinatario y produce el cuerpo de la carta.
    """
    from openai import OpenAI

    api_key = os.getenv("OPENAI_API_KEY")
    print(f"DEBUG OPENAI_API_KEY: {'SET (' + api_key[:10] + '...)' if api_key else 'NOT SET'}", flush=True)
    if not api_key:
        raise HTTPException(status_code=503, detail="API de IA no configurada (OPENAI_API_KEY)")

    # Obtener datos del contrato
    contrato = db.query(Contrato).filter(Contrato.id == request.contrato_id).first()
    if not contrato:
        raise HTTPException(status_code=404, detail="Contrato no encontrado")

    # Número de carta correlativo
    correlativo, numero_carta = _obtener_siguiente_numero_carta(db)

    # Fecha en formato peruano
    meses = ["enero","febrero","marzo","abril","mayo","junio",
             "julio","agosto","septiembre","octubre","noviembre","diciembre"]
    hoy = datetime.now()
    fecha_texto = f"Lima, {hoy.day} de {meses[hoy.month-1]} de {hoy.year}"

    # Destinatario desde el contrato
    destinatario_nombre = contrato.nombre_representante or contrato.contratado or ""
    destinatario_cargo = contrato.cargo_representante or ""
    destinatario_institucion = contrato.contratado or ""

    # Referencia a plantilla si hay
    texto_plantilla = ""
    if request.plantilla_id:
        plantilla = db.query(PlantillaCarta).filter(PlantillaCarta.id == request.plantilla_id).first()
        if plantilla and plantilla.archivo_local:
            ruta = os.path.join(UPLOAD_DIR, plantilla.archivo_local)
            if os.path.exists(ruta):
                texto_plantilla = _leer_plantilla_docx(ruta)

    # Contexto del contrato para la IA
    contexto_contrato = f"""
Contrato N° {contrato.numero or 'S/N'}
Tipo: {contrato.tipo_contrato or ''}
Contratado: {contrato.contratado or ''}
Item: {contrato.item_contratado or ''}
Estado: {contrato.estado_ejecucion or ''}
"""

    prompt_sistema = """Eres un asistente experto en redacción de cartas institucionales formales peruanas.
Redactas cartas en nombre del NÚCLEO EJECUTOR PARA EL MANTENIMIENTO, ACONDICIONAMIENTO Y EQUIPAMIENTO DE COMISARÍAS - NEMAEC.
El firmante siempre es MIGUEL IVAN ALARCÓN PARCO, PRESIDENTE NEMAEC.
Responde ÚNICAMENTE con un JSON con esta estructura exacta (sin markdown, sin explicaciones):
{
  "referencias": "a) ...\nb) ...",
  "cuerpo": "Tengo el agrado de dirigirme a usted...",
  "cierre": "Sin otro particular, hago propicia la oportunidad para expresarle los sentimientos de mi especial consideración y estima."
}"""

    prompt_usuario = f"""Redacta una carta institucional formal con la siguiente información:

DATOS DEL CONTRATO:
{contexto_contrato}

ASUNTO: {request.asunto}

REFERENCIAS INDICADAS: {request.referencias or '(ninguna)'}

INSTRUCCIONES ADICIONALES: {request.instrucciones or '(ninguna)'}

{f'EJEMPLO DE ESTRUCTURA (carta de referencia):{chr(10)}{texto_plantilla[:1500]}' if texto_plantilla else ''}

Genera:
1. "referencias": las referencias formales del contrato en formato "a) ...\nb) ..."
2. "cuerpo": el cuerpo completo de la carta (desde "Tengo el agrado..." hasta antes del cierre), formal, conciso y con el tono institucional peruano correcto
3. "cierre": la frase de cierre (usar la estándar si no se indica otra)"""

    try:
        client = OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": prompt_sistema},
                {"role": "user", "content": prompt_usuario}
            ],
            temperature=0.3,
            max_tokens=1500,
        )
        contenido = response.choices[0].message.content.strip()
        # Limpiar posibles bloques markdown
        if contenido.startswith("```"):
            contenido = re.sub(r'^```[a-z]*\n?', '', contenido)
            contenido = re.sub(r'\n?```$', '', contenido)
        datos = json.loads(contenido)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generando carta con IA: {str(e)}")

    return GenerarCartaResponse(
        numero_carta=numero_carta,
        fecha_texto=fecha_texto,
        destinatario_nombre=destinatario_nombre.upper(),
        destinatario_cargo=destinatario_cargo,
        destinatario_institucion=destinatario_institucion.upper(),
        asunto=request.asunto.upper(),
        referencias=datos.get("referencias", ""),
        cuerpo=datos.get("cuerpo", ""),
        cierre=datos.get("cierre", "Sin otro particular, hago propicia la oportunidad para expresarle los sentimientos de mi especial consideración y estima."),
    )


def _construir_docx_buffer(request: ExportarCartaRequest, db: Session):
    """
    Genera el .docx de una carta y lo retorna como BytesIO.
    Función interna reutilizable por guardar_carta y exportar_carta_docx.
    """
    import io
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor, Cm as DocxCm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn, qn as _qn
        from docx.oxml import OxmlElement as _OxmlElement
    except ImportError:
        raise HTTPException(status_code=503, detail="python-docx no instalado.")

    def get_cfg(clave, default=""):
        row = db.query(ConfiguracionSistema).filter(ConfiguracionSistema.clave == clave).first()
        return row.valor if row and row.valor else default

    # ── Abrir membrete como base (si existe), si no documento en blanco ──
    cfg_membrete = db.query(ConfiguracionSistema).filter(
        ConfiguracionSistema.clave == "membrete_archivo"
    ).first()
    ruta_membrete = None
    if cfg_membrete and cfg_membrete.valor:
        ruta_membrete = os.path.join(UPLOAD_DIR, cfg_membrete.valor)
        if not os.path.exists(ruta_membrete):
            ruta_membrete = None

    if ruta_membrete:
        doc = Document(ruta_membrete)
        body = doc.element.body
        sect_pr = body.find(qn('w:sectPr'))
        for child in list(body):
            if child != sect_pr:
                body.remove(child)
    else:
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    estilo = doc.styles['Normal']
    estilo.font.name = 'Arial'
    estilo.font.size = Pt(11)

    def agregar_parrafo(texto, bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT,
                        space_before=0, space_after=6, color=None):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(texto)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = 'Arial'
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p

    agregar_parrafo(request.fecha_texto, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=8)
    agregar_parrafo(request.numero_carta, bold=True, space_after=12)
    agregar_parrafo("Señor:", space_after=0)
    agregar_parrafo(request.destinatario_nombre, bold=True, space_after=0)
    if request.destinatario_cargo:
        agregar_parrafo(request.destinatario_cargo, space_after=0)
    agregar_parrafo(request.destinatario_institucion, space_after=0)
    agregar_parrafo("Presente. –", space_after=12)

    p_asunto = doc.add_paragraph()
    p_asunto.paragraph_format.space_after = Pt(4)
    r1 = p_asunto.add_run("Asunto:\t"); r1.bold = True; r1.font.name = 'Arial'; r1.font.size = Pt(11)
    r2 = p_asunto.add_run(request.asunto); r2.bold = True; r2.font.name = 'Arial'; r2.font.size = Pt(11)

    if request.referencias and request.referencias.strip():
        p_ref = doc.add_paragraph()
        p_ref.paragraph_format.space_after = Pt(12)
        r_lbl = p_ref.add_run("Referencia:\t"); r_lbl.bold = True; r_lbl.font.name = 'Arial'; r_lbl.font.size = Pt(11)
        r_val = p_ref.add_run(request.referencias); r_val.font.name = 'Arial'; r_val.font.size = Pt(11)
    else:
        agregar_parrafo("", space_after=12)

    agregar_parrafo("De mi especial consideración.", space_after=8)

    for linea in request.cuerpo.split('\n'):
        if linea.strip():
            p = agregar_parrafo(linea.strip(), space_after=6)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p_cierre = agregar_parrafo(request.cierre, space_after=24)
    p_cierre.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    agregar_parrafo("Atentamente,", space_after=0)

    firma_imagen = get_cfg("firma_imagen")
    firma_nombre = get_cfg("firma_nombre", "").strip()
    firma_cargo  = get_cfg("firma_cargo", "").strip()

    if firma_imagen:
        ruta_firma = os.path.join(UPLOAD_DIR, firma_imagen)
        if os.path.exists(ruta_firma):
            # ── Recortar espacio en blanco de la imagen de firma ──
            ruta_firma_uso = ruta_firma
            ruta_firma_tmp = None
            try:
                from PIL import Image as _PILImage
                import tempfile as _tempfile
                img = _PILImage.open(ruta_firma)
                img_gray = img.convert('L')
                # Máscara: píxeles no blancos (firma es oscura, fondo es blanco)
                mask = img_gray.point(lambda x: 0 if x > 240 else 255)
                bbox = mask.getbbox()
                if bbox:
                    pad = 10  # px de margen alrededor de la firma
                    left  = max(0, bbox[0] - pad)
                    top   = max(0, bbox[1] - pad)
                    right = min(img.width,  bbox[2] + pad)
                    bottom= min(img.height, bbox[3] + pad)
                    img_recortada = img.crop((left, top, right, bottom))
                    tmp = _tempfile.NamedTemporaryFile(suffix='.png', delete=False)
                    ruta_firma_tmp = tmp.name
                    tmp.close()
                    img_recortada.save(ruta_firma_tmp, 'PNG')
                    ruta_firma_uso = ruta_firma_tmp
            except Exception as _e:
                print(f"Advertencia al recortar firma: {_e}")

            # ── Imagen de firma (derecha, compacta, sobre la línea) ──
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_img.paragraph_format.space_before = Pt(6)
            p_img.paragraph_format.space_after = Pt(0)
            r_img = p_img.add_run()
            r_img.add_picture(ruta_firma_uso, height=DocxCm(2.0))

            if ruta_firma_tmp:
                try:
                    os.unlink(ruta_firma_tmp)
                except Exception:
                    pass
        else:
            agregar_parrafo("", space_before=36, space_after=0)
    else:
        agregar_parrafo("", space_before=36, space_after=0)

    # ── Línea + nombre + cargo (siempre a la derecha) ──
    agregar_parrafo("____________________________", align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=0)
    if firma_nombre:
        agregar_parrafo(firma_nombre, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    if firma_cargo:
        agregar_parrafo(firma_cargo, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _docx_a_pdf(ruta_docx: str, ruta_pdf: str) -> bool:
    """
    Convierte un .docx a PDF. Intenta con docx2pdf (Word en Windows),
    luego con LibreOffice como fallback. Retorna True si tuvo éxito.
    """
    # Intento 1: docx2pdf (usa Microsoft Word via COM en Windows)
    try:
        from docx2pdf import convert
        convert(ruta_docx, ruta_pdf)
        if os.path.exists(ruta_pdf) and os.path.getsize(ruta_pdf) > 0:
            return True
    except Exception as e:
        print(f"docx2pdf falló: {e}")

    # Intento 2: LibreOffice headless
    try:
        import subprocess
        out_dir = os.path.dirname(ruta_pdf)
        result = subprocess.run(
            ['soffice', '--headless', '--convert-to', 'pdf', '--outdir', out_dir, ruta_docx],
            capture_output=True, timeout=60
        )
        # LibreOffice genera el PDF con el mismo nombre base del .docx
        nombre_base = os.path.splitext(os.path.basename(ruta_docx))[0]
        pdf_generado = os.path.join(out_dir, nombre_base + '.pdf')
        if os.path.exists(pdf_generado):
            if pdf_generado != ruta_pdf:
                os.rename(pdf_generado, ruta_pdf)
            return os.path.exists(ruta_pdf) and os.path.getsize(ruta_pdf) > 0
    except Exception as e:
        print(f"LibreOffice también falló: {e}")

    return False


@app.post("/api/guardar-carta")
def guardar_carta(
    request: ExportarCartaRequest,
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    """
    Guarda la carta en el sistema:
    - Verifica que el número no exista ya para ese año
    - Si hay conflicto retorna 409 con el siguiente número disponible
    - Si OK: genera .docx + PDF, guarda en uploads, registra en documentos y expediente
    """
    def get_cfg(clave, default=""):
        row = db.query(ConfiguracionSistema).filter(ConfiguracionSistema.clave == clave).first()
        return row.valor if row and row.valor else default

    m = re.search(r'(\d+)-(\d{4})', request.numero_carta)
    if not m:
        raise HTTPException(status_code=400, detail="Formato de número de carta no reconocido")
    correlativo = int(m.group(1))
    anio = int(m.group(2))

    existe = db.query(CartaGenerada).filter(
        CartaGenerada.numero_correlativo == correlativo,
        CartaGenerada.anio == anio
    ).first()

    if existe:
        ultima = (
            db.query(CartaGenerada)
            .filter(CartaGenerada.anio == anio)
            .order_by(CartaGenerada.numero_correlativo.desc())
            .first()
        )
        siguiente = (ultima.numero_correlativo + 1) if ultima else 1
        digitos = int(get_cfg("carta_digitos", "6"))
        sufijo = get_cfg("carta_sufijo", "").strip()
        correlativo_str = str(siguiente).zfill(digitos)
        numero_sugerido = f"Carta N° {correlativo_str}-{anio}-{sufijo}" if sufijo else f"Carta N° {correlativo_str}-{anio}"
        raise HTTPException(
            status_code=409,
            detail=f"El número '{request.numero_carta}' ya existe. Número disponible: {numero_sugerido}"
        )

    contrato_id = request.contrato_id

    # Registrar en cartas_generadas
    nueva_carta = CartaGenerada(
        numero_correlativo=correlativo,
        anio=anio,
        numero_completo=request.numero_carta,
        contrato_id=contrato_id,
        asunto=request.asunto,
    )
    db.add(nueva_carta)

    # Registrar en documentos (sin archivos aún)
    nuevo_doc = Documento(
        tipo_documento='carta',
        direccion='enviado',
        numero=request.numero_carta,
        fecha=datetime.now(),
        destinatario=f"{request.destinatario_nombre} - {request.destinatario_institucion}",
        asunto=request.asunto,
        titulo=request.asunto,
        resumen=request.cuerpo[:500] if request.cuerpo else "",
        correlativo_oficio=correlativo,
        anio_oficio=anio,
        estado='borrador',
    )
    db.add(nuevo_doc)
    db.flush()  # Obtener nuevo_doc.id

    # ── Generar y guardar .docx ──
    nombre_base = re.sub(r'[^\w\-]', '_', request.numero_carta.replace(' ', '_').replace('°', ''))
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    nombre_docx = f"carta_{nuevo_doc.id}_{timestamp}_{nombre_base}.docx"
    nombre_pdf  = f"carta_{nuevo_doc.id}_{timestamp}_{nombre_base}.pdf"
    ruta_docx = os.path.join(UPLOAD_DIR, nombre_docx)
    ruta_pdf  = os.path.join(UPLOAD_DIR, nombre_pdf)

    try:
        buffer_docx = _construir_docx_buffer(request, db)
        with open(ruta_docx, 'wb') as f:
            f.write(buffer_docx.getvalue())
        nuevo_doc.archivo_docx = nombre_docx

        # Convertir a PDF
        if _docx_a_pdf(ruta_docx, ruta_pdf):
            nuevo_doc.archivo_local = nombre_pdf
        else:
            print(f"Advertencia: no se pudo convertir a PDF para carta {request.numero_carta}")
    except Exception as e:
        print(f"Error generando archivos de carta: {e}")
        # Continuar sin archivos — la carta queda registrada igualmente

    # Registrar en expediente del contrato si hay contrato_id
    if contrato_id:
        contrato = db.query(Contrato).filter(Contrato.id == contrato_id).first()
        if contrato:
            expediente = ExpedienteContrato(
                contrato_id=contrato_id,
                tipo_doc='Carta Enviada',
                numero=request.numero_carta,
                fecha=datetime.now(),
                asunto=request.asunto,
                archivo_local=nombre_pdf if nuevo_doc.archivo_local else None,
            )
            db.add(expediente)

    db.commit()

    return {
        "ok": True,
        "numero_carta": request.numero_carta,
        "documento_id": nuevo_doc.id,
        "tiene_pdf": bool(nuevo_doc.archivo_local),
        "tiene_docx": bool(nuevo_doc.archivo_docx),
        "mensaje": f"Carta '{request.numero_carta}' guardada en el sistema"
    }


@app.post("/api/exportar-carta")
def exportar_carta_docx(
    request: ExportarCartaRequest,
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    """
    Genera y descarga el .docx de la carta (sin guardar en el sistema).
    Usa _construir_docx_buffer internamente.
    """
    try:
        buffer = _construir_docx_buffer(request, db)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generando carta: {str(e)}")

    num_limpio = re.sub(r'[^\w\-]', '_', request.numero_carta.replace(' ', '_').replace('°', ''))
    nombre_descarga = f"{num_limpio}.docx"

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{nombre_descarga}"'}
    )


# ============================================
# SEGUIMIENTO LIQUIDACIÓN
# ============================================

CAMPOS_SIONO = {
    'acta_revisada', 'acta_remitida_ugpe',
    'mod_presentado_ne', 'mod_revisado_aprobado', 'mod_remitido_ugpe',
    'amp_presentado_ne', 'amp_revisado_aprobado', 'amp_adenda_firmada',
    'dossier_presentado_ne', 'dossier_revisado_aprobado', 'dossier_remitido_ugpe', 'dossier_remitido_pago',
    'liq_presentado_ne', 'liq_revisado_aprobado', 'liq_remitido_pago',
}

def _color(hex_color):
    return PatternFill("solid", fgColor=hex_color)

def _border():
    thin = Side(style='thin', color='000000')
    return Border(left=thin, right=thin, top=thin, bottom=thin)

def _font(bold=False, color="000000", size=9):
    return Font(bold=bold, color=color, size=size)

def _align(horizontal="center", wrap=True):
    return Alignment(horizontal=horizontal, vertical="center", wrap_text=wrap)

@app.get("/api/seguimiento/exportar-excel")
def exportar_seguimiento_excel(db: Session = Depends(get_db)):
    """Exporta la tabla de seguimiento como Excel. SI→✔ NO→✘ con formato condicional."""
    from openpyxl.styles.differential import DifferentialStyle
    from openpyxl.formatting.rule import Rule

    filas = db.query(SeguimientoComisaria).order_by(SeguimientoComisaria.numero).all()

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Seguimiento Liquidación"
    ws.sheet_view.showGridLines = False

    # Paleta gris – mismos tonos que el tema "gris" de la web
    C_TITLE   = "1E293B"   # slate-800  – barra de título
    C_GRP     = "334155"   # slate-700  – encabezados de sección
    C_SUB     = "475569"   # slate-600  – sub-encabezados
    C_HDR_TXT = "FFFFFF"
    C_TH_BG   = "F1F5F9"   # slate-100 – filas de encabezado (texto oscuro)
    C_TH_TXT  = "1E293B"
    C_ALT     = "F8FAFC"   # slate-50  – filas pares
    C_NA      = "E2E8F0"   # slate-200
    C_NA_TXT  = "64748B"   # slate-500
    # Celdas ✔/✘
    C_SI_BG   = "DCFCE7"   # green-100
    C_SI_TXT  = "166534"   # green-800
    C_NO_BG   = "FEE2E2"   # red-100
    C_NO_TXT  = "991B1B"   # red-800

    ICO_SI = "✔"
    ICO_NO = "✘"

    brd = _border()

    # ── FILA 1: Título ──────────────────────────────────────────────
    ws.merge_cells("A1:X1")
    c = ws["A1"]
    c.value = "SEGUIMIENTO AL PROCESO DE LIQUIDACIÓN - MANTENIMIENTO Y ACONDICIONAMIENTO DE COMISARÍAS"
    c.fill = _color(C_TITLE)
    c.font = _font(bold=True, color=C_HDR_TXT, size=11)
    c.alignment = _align("center")
    c.border = brd
    ws.row_dimensions[1].height = 22

    # ── FILA 2: Encabezados de grupo ────────────────────────────────
    grupos = [
        ("A2","A3","N°",              C_TITLE),
        ("B2","B3","COMISARÍA PNP",  C_TITLE),
        ("C2","D2","AVANCE DE EJECUCIÓN", C_TITLE),
        ("E2","E3","FECHA FINAL\nEJECUCIÓN CONTRACTUAL", C_TITLE),
        ("F2","H2","1. ACTA DE CONFORMIDAD\nDE EJECUCIÓN Y RECEPCIÓN FÍSICA", C_GRP),
        ("I2","K2","2. INFORME DE MODIFICACIÓN\nDE PARTIDAS (UGPE)", C_GRP),
        ("L2","N2","3. INFORME DE\nAMPLIACIÓN DE PLAZO", C_GRP),
        ("O2","S2","4. INFORME DE CULMINACIÓN Y\nENTREGA DE OBRA (DOSSIER)", C_GRP),
        ("T2","V2","5. INFORME DE LIQUIDACIÓN\n(FINAL)", C_GRP),
        ("W2","W3","OBSERVACIONES", C_TITLE),
    ]
    for start, end, label, color in grupos:
        if start != end:
            ws.merge_cells(f"{start}:{end}")
        cel = ws[start]
        cel.value = label
        cel.fill = _color(color)
        cel.font = _font(bold=True, color=C_HDR_TXT, size=8)
        cel.alignment = _align("center")
        cel.border = brd
    ws.row_dimensions[2].height = 30

    # ── FILA 3: Sub-encabezados ─────────────────────────────────────
    ws.merge_cells("A2:A3")
    ws.merge_cells("B2:B3")
    ws.merge_cells("E2:E3")
    ws.merge_cells("W2:W3")

    sub_hdrs = [
        ("C3", "PROGRAMADO"),
        ("D3", "FÍSICO"),
        ("F3", "FECHA FIRMA\nACTA"),
        ("G3", "REVISADA Y\nAPROBADA"),
        ("H3", "REMITIDA\nA UGPE"),
        ("I3", "PRESENTADO\nAL NE"),
        ("J3", "REVISADO Y\nAPROBADO"),
        ("K3", "REMITIDO\nA UGPE"),
        ("L3", "PRESENTADO\nAL NE"),
        ("M3", "REVISADO Y\nAPROBADO"),
        ("N3", "ADENDA\nFIRMADA"),
        ("O3", "PRESENTADO\nAL NE"),
        ("P3", "REVISADO Y\nAPROBADO"),
        ("Q3", "REMITIDO\nA UGPE"),
        ("R3", "REMITIDO\nPARA PAGO"),
        ("S3", "MONTO\nPAGADO (S/)"),
        ("T3", "PRESENTADO\nAL NE"),
        ("U3", "REVISADO Y\nAPROBADO"),
        ("V3", "REMITIDO\nPARA PAGO"),
    ]
    for cell_ref, label in sub_hdrs:
        c = ws[cell_ref]
        c.value = label
        c.fill = _color(C_SUB)
        c.font = _font(bold=True, color=C_HDR_TXT, size=8)
        c.alignment = _align("center")
        c.border = brd
    ws.row_dimensions[3].height = 30

    # ── FILAS DE DATOS ───────────────────────────────────────────────
    campos_siono = [
        'acta_revisada','acta_remitida_ugpe',
        'mod_presentado_ne','mod_revisado_aprobado','mod_remitido_ugpe',
        'amp_presentado_ne','amp_revisado_aprobado','amp_adenda_firmada',
        'dossier_presentado_ne','dossier_revisado_aprobado','dossier_remitido_ugpe','dossier_remitido_pago',
        'liq_presentado_ne','liq_revisado_aprobado','liq_remitido_pago',
    ]
    # G(7)..R(18) + saltar S(19)=monto → T(20)..V(22)
    col_map = {campo: idx for idx, campo in enumerate(campos_siono, start=7)}
    for campo, col in list(col_map.items()):
        if col >= 19:
            col_map[campo] = col + 1

    data_first_row = 4
    for r_idx, row in enumerate(filas, start=data_first_row):
        bg = None if r_idx % 2 == 0 else C_ALT

        def cell(col_num, r=r_idx):
            return ws.cell(row=r, column=col_num)

        cell(1).value = row.numero
        cell(1).font = _font(bold=True)
        cell(2).value = row.comisaria
        cell(2).font = _font(bold=False)
        cell(2).alignment = _align("left", wrap=False)
        cell(3).value = row.avance_programado
        cell(3).number_format = "0%"
        cell(4).value = row.avance_fisico
        cell(4).number_format = "0.00%"
        if row.fecha_fin_contractual:
            cell(5).value = row.fecha_fin_contractual
            cell(5).number_format = "DD/MM/YYYY"
        if row.acta_fecha_firma:
            cell(6).value = row.acta_fecha_firma
            cell(6).number_format = "DD/MM/YYYY"

        # SI → ✔ (verde), NO → ✘ (rojo), NA/- → gris
        for campo, col_num in col_map.items():
            val = getattr(row, campo) or ''
            c = cell(col_num)
            c.alignment = _align("center")
            if val == 'SI':
                c.value = ICO_SI
                c.fill = _color(C_SI_BG)
                c.font = _font(bold=True, color=C_SI_TXT, size=11)
            elif val == 'NO':
                c.value = ICO_NO
                c.fill = _color(C_NO_BG)
                c.font = _font(bold=True, color=C_NO_TXT, size=11)
            elif val in ('NA', '-', ''):
                c.value = '–' if not val else val
                c.fill = _color(C_NA)
                c.font = _font(color=C_NA_TXT)

        # Monto pagado (col T = 20)
        if row.dossier_monto_pagado is not None:
            cell(20).value = row.dossier_monto_pagado
            cell(20).number_format = '#,##0.00'
            cell(20).alignment = _align("right")

        cell(24).value = row.observaciones or ''
        cell(24).alignment = _align("left")

        # Borde y fondo alternado
        for col_num in range(1, 25):
            c = cell(col_num)
            c.border = brd
            try:
                rgb = c.fill.fgColor.rgb
            except Exception:
                rgb = "00000000"
            if bg and rgb in ("00000000", "00FFFFFF", "FFFFFFFF"):
                c.fill = _color(bg)
            if c.alignment.horizontal is None:
                c.alignment = _align("center")

        ws.row_dimensions[r_idx].height = 15

    last_data_row = data_first_row + len(filas) - 1

    # ── ANCHOS DE COLUMNA ────────────────────────────────────────────
    col_widths = {1:4, 2:22, 3:7, 4:8.43, 5:11, 6:11,
                  7:7,8:7,9:7,10:7,11:7,12:7,13:7,14:7,15:7,
                  16:7,17:7,18:7,19:7,20:13,21:7,22:7,23:7,24:22}
    for col_num, width in col_widths.items():
        ws.column_dimensions[get_column_letter(col_num)].width = width

    # ── FILA TOTALES ─────────────────────────────────────────────────
    total_row = last_data_row + 1
    ws.merge_cells(f"A{total_row}:B{total_row}")
    ws.cell(total_row, 1).value = "TOTAL / PROMEDIO"
    ws.cell(total_row, 1).font = _font(bold=True, size=9)
    ws.cell(total_row, 1).alignment = _align("right")
    ws.cell(total_row, 3).value = f"=AVERAGE(C{data_first_row}:C{last_data_row})"
    ws.cell(total_row, 3).number_format = "0%"
    ws.cell(total_row, 4).value = f"=AVERAGE(D{data_first_row}:D{last_data_row})"
    ws.cell(total_row, 4).number_format = "0.00%"
    # Total monto pagado (col T=20)
    ws.cell(total_row, 20).value = f"=SUM(T{data_first_row}:T{last_data_row})"
    ws.cell(total_row, 20).number_format = '#,##0.00'
    ws.cell(total_row, 20).alignment = _align("right")
    ws.cell(total_row, 20).font = _font(bold=True, color="145f2e", size=10)
    for col_num in range(1, 25):
        c = ws.cell(total_row, col_num)
        c.border = brd
        c.fill = _color("E2E8F0")
    ws.row_dimensions[total_row].height = 16

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    fecha = datetime.now().strftime("%d.%m.%Y")
    nombre = f"Seguimiento Liquidacion Comisarias ({fecha}).xlsx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{nombre}"'}
    )


@app.get("/api/seguimiento", response_model=List[SeguimientoComisariaResponse])
def get_seguimiento(db: Session = Depends(get_db)):
    """Retorna todas las filas de seguimiento. Público (sin autenticación)."""
    return db.query(SeguimientoComisaria).order_by(SeguimientoComisaria.numero).all()


@app.put("/api/seguimiento/{comisaria_id}/celda")
def actualizar_celda(
    comisaria_id: int,
    request: ActualizarCeldaRequest,
    db: Session = Depends(get_db),
    payload: dict = Depends(verificar_admin)
):
    """Actualiza el valor de un campo en una comisaría. Requiere autenticación."""
    comisaria = db.query(SeguimientoComisaria).filter(SeguimientoComisaria.id == comisaria_id).first()
    if not comisaria:
        raise HTTPException(status_code=404, detail="Comisaría no encontrada")

    campo = request.campo
    if not hasattr(comisaria, campo):
        raise HTTPException(status_code=400, detail=f"Campo '{campo}' no existe")

    valor_anterior = getattr(comisaria, campo)

    # Convertir el valor según el tipo del campo
    CAMPOS_FECHA = {'fecha_fin_contractual', 'acta_fecha_firma'}
    CAMPOS_FLOAT = {'avance_fisico', 'avance_programado', 'dossier_monto_pagado'}
    CAMPOS_BOOL = {'dossier_monto_merge', 'amp_merge'}
    valor_convertido = request.valor
    if request.valor is not None and request.valor != '':
        if campo in CAMPOS_FECHA:
            try:
                valor_convertido = datetime.strptime(request.valor[:10], '%Y-%m-%d')
            except ValueError:
                raise HTTPException(status_code=400, detail="Formato de fecha inválido (esperado YYYY-MM-DD)")
        elif campo in CAMPOS_FLOAT:
            try:
                valor_convertido = float(request.valor)
            except ValueError:
                raise HTTPException(status_code=400, detail="Valor numérico inválido")
        elif campo in CAMPOS_BOOL:
            valor_convertido = request.valor.lower() in ('true', '1', 'si', 'yes')
    elif request.valor == '':
        valor_convertido = None

    setattr(comisaria, campo, valor_convertido)
    comisaria.updated_at = datetime.now()

    nombre_usuario = payload.get("sub") or payload.get("username", "desconocido")

    # Si cambia a SI y hay detalle, registrarlo
    if request.valor == 'SI' and campo in CAMPOS_SIONO:
        if request.observacion or request.enlace:
            detalle = SeguimientoCeldaDetalle(
                comisaria_id=comisaria_id,
                campo=campo,
                observacion=request.observacion,
                enlace=request.enlace,
                usuario=nombre_usuario,
                fecha_actualizacion=datetime.now()
            )
            db.add(detalle)

    db.commit()
    db.refresh(comisaria)
    return {"ok": True, "valor_anterior": valor_anterior, "valor_nuevo": request.valor}


@app.post("/api/seguimiento/{comisaria_id}/celda/{campo}/archivo")
async def subir_archivo_celda(
    comisaria_id: int,
    campo: str,
    archivo: UploadFile = File(...),
    observacion: Optional[str] = None,
    enlace: Optional[str] = None,
    db: Session = Depends(get_db),
    payload: dict = Depends(verificar_admin)
):
    """Sube un archivo adjunto al detalle de una celda. Requiere autenticación."""
    comisaria = db.query(SeguimientoComisaria).filter(SeguimientoComisaria.id == comisaria_id).first()
    if not comisaria:
        raise HTTPException(status_code=404, detail="Comisaría no encontrada")

    ext = os.path.splitext(archivo.filename)[1]
    nombre_archivo = f"seg_{comisaria_id}_{campo}_{datetime.now().strftime('%Y%m%d%H%M%S')}{ext}"
    ruta = os.path.join(UPLOAD_DIR, nombre_archivo)
    with open(ruta, "wb") as f:
        f.write(await archivo.read())

    nombre_usuario = payload.get("sub") or payload.get("username", "desconocido")
    detalle = SeguimientoCeldaDetalle(
        comisaria_id=comisaria_id,
        campo=campo,
        observacion=observacion,
        enlace=enlace,
        archivo_local=nombre_archivo,
        archivo_nombre=archivo.filename,
        usuario=nombre_usuario,
        fecha_actualizacion=datetime.now()
    )
    db.add(detalle)
    db.commit()
    return {"ok": True, "archivo": nombre_archivo, "ruta": f"/uploads/{nombre_archivo}"}


# ============================================
# SERVIR FRONTEND
# ============================================

# Directorio del frontend
FRONTEND_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "frontend")

# Montar archivos estáticos del frontend (CSS, JS)
if os.path.exists(FRONTEND_DIR):
    app.mount("/css", StaticFiles(directory=os.path.join(FRONTEND_DIR, "css")), name="css")
    app.mount("/js", StaticFiles(directory=os.path.join(FRONTEND_DIR, "js")), name="js")

# ============================================
# KAIZEN — REGISTRO DE MEJORA
# ============================================

_PROMPT_KAIZEN = """Eres un asistente de redacción para registros Kaizen de NEMAEC \
(núcleo ejecutor FONCODES, mantenimiento y equipamiento de comisarías).

Si el contexto incluye "Puesto del usuario:", usa vocabulario del rol:
- MONITOR DE OBRA → obra, partidas, contratista, cuaderno de obra, valorizaciones.
- ADMINISTRADOR / ASISTENTE ADMINISTRATIVO → expedientes, trámites, plazos, firmas.
- COORDINADOR / ASISTENTE DE EQUIPAMIENTO → actas, recepción, inventario, proveedor.
- COORDINADOR GENERAL / PRESIDENTE → UGPE, FONCODES, informes, coordinación.
- TESORERO → pagos, transferencias, presupuesto, rendición.
- ESPECIALISTA TIC → sistemas, servidores, acceso, datos, red.
- SECRETARIO → correspondencia, archivo, oficios, plazos de respuesta.

REGLAS DE COMPORTAMIENTO:
1. TURNO 1 (sin historial previo): Si el texto es muy vago o vacío, haz UNA sola pregunta \
   concreta. Si ya tiene contenido útil, ve directo a redactar.
2. TURNO 2 en adelante (hay historial): NUNCA hagas otra pregunta. Con todo lo que el usuario \
   te dijo, redacta la versión final y pregunta solo: \
   "¿Estás de acuerdo con esta formulación? «[texto concreto y completo]»"
3. Si el usuario confirma (dice sí, ok, bien, correcto, de acuerdo, perfecto): \
   responde únicamente "¡Perfecto! Pulsa el botón amarillo para aplicarlo al formulario."
4. Máximo 3 oraciones en total por respuesta.
5. El texto entre «» debe ser autocontenido y listo para pegar en un formulario oficial.
6. No repitas preguntas que ya hiciste en el historial."""

_GUIAS_BLOQUE = {
    "problema":    "El usuario describe un problema del proyecto. Ayúdalo a ser más específico sobre: qué ocurrió exactamente, cuándo, dónde y cuánto afectó.",
    "impacto":     "El usuario describe el impacto. Ayúdalo a cuantificar: días de retraso, costo adicional, quién se vio afectado.",
    "causa":       "El usuario responde a un '¿Por qué?'. Detecta si es una consecuencia (no una causa raíz) y empújalo a ir más profundo.",
    "solucion":    "El usuario propone una solución. Ayúdalo a hacerla concreta: ¿qué acción específica, quién, cuándo exactamente?",
    "aprendizaje": "El usuario escribe un aprendizaje. Ayúdalo a convertirlo en una regla o estándar aplicable: 'Siempre... antes de...'",
}

@app.get("/api/mejoras", response_model=List[RegistroMejoraResponse])
def listar_mejoras(
    estado: Optional[str] = Query(None),
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    q = db.query(RegistroMejora)
    if estado:
        q = q.filter(RegistroMejora.estado == estado)
    return q.order_by(RegistroMejora.created_at.desc()).all()

@app.post("/api/mejoras", response_model=RegistroMejoraResponse, status_code=201)
def crear_mejora(
    datos: RegistroMejoraCreate,
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    registro = RegistroMejora(**datos.model_dump(), usuario=admin.get("nombre", admin.get("sub", "usuario")))
    db.add(registro)
    db.commit()
    db.refresh(registro)
    return registro

@app.post("/api/mejoras/asistir", response_model=AsistirMejoraResponse)
def asistir_mejora(
    req: AsistirMejoraRequest,
    admin: dict = Depends(verificar_admin)
):
    """Asistencia IA: SOLO hace preguntas y pide precisión. Nunca genera contenido."""
    api_key = os.getenv("OPENAI_API_KEY")

    guia = _GUIAS_BLOQUE.get(req.bloque, "Ayuda al usuario a ser más específico.")
    acciones = {
        "mejorar":    "El usuario quiere mejorar la redacción. Señala qué parte es imprecisa y pregunta algo concreto.",
        "especificar":"El usuario quiere ser más específico. Pregunta por detalles: qué, cuándo, cuánto, dónde.",
        "ayudar":     "El usuario no sabe cómo continuar. Hazle 1-2 preguntas guía muy concretas.",
        "convertir":  "El usuario quiere convertir su texto en algo accionable. Pregunta qué acción concreta implica.",
        "preguntar":  "Formula la siguiente pregunta del proceso 5 Porqués basándote en lo que el usuario ya escribió.",
        "validar":    "El usuario quiere saber si llegó a la causa raíz. Evalúa si la cadena de porqués es profunda o superficial.",
        "reformular": "El usuario quiere reformular su texto. Señala qué parte mejorar y pregunta algo concreto.",
    }
    instruccion_accion = acciones.get(req.accion, "Ayuda al usuario a ser más preciso.")
    contexto_extra = f"\nContexto adicional: {req.contexto}" if req.contexto else ""

    prompt_usuario = f"""Bloque: {req.bloque}
Instrucción: {guia}
Acción solicitada: {instruccion_accion}{contexto_extra}

Texto del usuario:
\"\"\"{req.texto}\"\"\"

Responde SOLO con JSON:
{{"respuesta": "...", "tipo": "pregunta|sugerencia|reformulacion"}}"""

    if not api_key:
        puesto = ""
        if req.contexto:
            for linea in req.contexto.splitlines():
                if linea.startswith("Puesto del usuario:"):
                    puesto = linea.split(":", 1)[1].strip().upper()
                    break

        mocks_por_puesto = {
            "MONITOR DE OBRA": {
                "problema":    ("¿En qué partida o actividad ocurrió? Con ese dato puedo ayudarte a redactarlo. Si ya lo tienes escrito, puedes usar: «Durante la ejecución de [partida], se detectó que [describe el problema], lo que impidió continuar con el avance programado.»", "sugerencia"),
                "impacto":     ("Puedes usar: «El problema generó un retraso de [X] días en el cronograma de ejecución, afectando la fecha de entrega prevista y obligando a levantar una anotación en el cuaderno de obra.»", "sugerencia"),
                "causa":       ("Puedes usar: «El contratista no contaba con [material/recurso] al momento de iniciar la actividad porque no gestionó el abastecimiento con anticipación suficiente.»", "sugerencia"),
                "solucion":    ("Puedes usar: «Se emitió una anotación en el cuaderno de obra exigiendo al contratista presentar un plan de recuperación en 48 horas, con responsable: [nombre/cargo].»", "sugerencia"),
                "aprendizaje": ("Puedes usar: «Antes de iniciar cada frente de trabajo, el monitor debe verificar que el contratista cuenta con los materiales y recursos necesarios para al menos 5 días de avance.»", "sugerencia"),
            },
            "ADMINISTRADOR": {
                "problema":    ("Puedes usar: «Se detectó que el expediente de [trámite] fue devuelto por [área/entidad] debido a [motivo], generando un retraso de [X] días en el proceso.»", "sugerencia"),
                "impacto":     ("Puedes usar: «El retraso bloqueó el procesamiento del pago/trámite correspondiente y generó una observación formal que requirió corrección y reenvío del expediente.»", "sugerencia"),
                "causa":       ("Puedes usar: «No existía un checklist de requisitos previo al envío del expediente, lo que permitió que se remitiera con documentos incompletos o sin firmas requeridas.»", "sugerencia"),
                "solucion":    ("Puedes usar: «Se implementó un checklist de verificación obligatorio antes de remitir cualquier expediente, con revisión del [cargo responsable] previo al despacho.»", "sugerencia"),
                "aprendizaje": ("Puedes usar: «Todo expediente debe pasar por una revisión de checklist antes de ser firmado y remitido, asegurando que todos los requisitos documentarios estén completos.»", "sugerencia"),
            },
            "ESPECIALISTA TIC": {
                "problema":    ("Puedes usar: «El sistema [nombre] estuvo inaccesible durante [X] horas debido a [causa técnica], impidiendo que [N] usuarios realizaran sus operaciones habituales.»", "sugerencia"),
                "impacto":     ("Puedes usar: «La interrupción bloqueó el registro de [tipo de operación] durante [X] horas, generando retrasos en los plazos de respuesta y afectando la operación del área.»", "sugerencia"),
                "causa":       ("Puedes usar: «La intervención no contaba con un plan de rollback aprobado ni con una ventana de mantenimiento coordinada, lo que impidió revertir el cambio rápidamente.»", "sugerencia"),
                "solucion":    ("Puedes usar: «Se estableció que toda intervención en servidores de producción debe realizarse en ventana de mantenimiento (viernes 6-8pm), con plan de rollback documentado y aprobado previamente.»", "sugerencia"),
                "aprendizaje": ("Puedes usar: «Antes de aplicar cualquier cambio en producción, se debe tener un plan de rollback probado y una ventana de mantenimiento comunicada con al menos 48 horas de anticipación.»", "sugerencia"),
            },
            "COORDINADOR DE EQUIPAMIENTO": {
                "problema":    ("Puedes usar: «Durante la recepción de [tipo de bien] en la comisaría [nombre], se detectó que [N] de [total] unidades presentaban [defecto], impidiendo su instalación inmediata.»", "sugerencia"),
                "impacto":     ("Puedes usar: «La observación retrasó la entrega final en [X] días hábiles y requirió levantar un acta de observación formal al proveedor para gestionar el reemplazo.»", "sugerencia"),
                "causa":       ("Puedes usar: «No se realizó una verificación física detallada al momento del desembalaje porque el protocolo de recepción no incluía ese paso como obligatorio.»", "sugerencia"),
                "solucion":    ("Puedes usar: «Se coordinó el reemplazo con el proveedor en garantía y se actualizó el protocolo de recepción incluyendo verificación fotográfica obligatoria antes de firmar el acta.»", "sugerencia"),
                "aprendizaje": ("Puedes usar: «Toda recepción de bienes debe incluir inspección visual ítem por ítem con registro fotográfico antes de firmar el acta de conformidad.»", "sugerencia"),
            },
        }
        mocks_genericos = {
            "problema":    ("Puedes usar: «Se identificó que [describe el problema] durante [momento/proceso], lo que afectó [área/actividad] por aproximadamente [duración o magnitud].»", "sugerencia"),
            "impacto":     ("Puedes usar: «El problema generó un retraso de [X] días/horas en [proceso afectado], requiriendo [acción correctiva inmediata].»", "sugerencia"),
            "causa":       ("Puedes usar: «La causa raíz fue la ausencia de [proceso/control/protocolo] que permitió que [situación problemática] ocurriera sin ser detectada a tiempo.»", "sugerencia"),
            "solucion":    ("Puedes usar: «Se implementó [acción concreta] a cargo de [responsable], con efecto a partir de [momento], para evitar que el problema se repita.»", "sugerencia"),
            "aprendizaje": ("Puedes usar: «Antes de [actividad crítica], siempre se debe verificar [control o condición] para evitar [consecuencia negativa].»", "sugerencia"),
        }
        mocks = mocks_por_puesto.get(puesto, mocks_genericos)
        txt, tipo = mocks.get(req.bloque, ("¿Puedes ser más específico sobre lo que describes?", "pregunta"))
        return AsistirMejoraResponse(respuesta=txt, tipo=tipo)

    from openai import OpenAI
    client = OpenAI(api_key=api_key)
    try:
        # Detectar turno: si hay historial con mensajes del usuario, estamos en turno 2+
        turno = 0
        if req.historial:
            turno = sum(1 for m in req.historial if m.get("role") == "user")

        if turno >= 1:
            prompt_usuario += "\n\n[INSTRUCCIÓN INTERNA: Es el turno 2+. NO hagas más preguntas. " \
                "Sintetiza TODO lo conversado y produce la formulación final entre «».]"

        messages = [
            {"role": "system", "content": _PROMPT_KAIZEN},
            {"role": "user",   "content": prompt_usuario},
        ]
        if req.historial:
            messages.extend(req.historial)
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=messages,
            temperature=0.3,
            max_tokens=300,
        )
        raw = resp.choices[0].message.content.strip()
        if raw.startswith("```"):
            raw = re.sub(r'^```[a-z]*\n?', '', raw)
            raw = re.sub(r'\n?```$', '', raw)
        datos = json.loads(raw)
        return AsistirMejoraResponse(
            respuesta=datos.get("respuesta", ""),
            tipo=datos.get("tipo", "pregunta")
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error IA: {str(e)}")

@app.put("/api/mejoras/{rid}", response_model=RegistroMejoraResponse)
def actualizar_mejora(
    rid: int,
    datos: RegistroMejoraUpdate,
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    registro = db.query(RegistroMejora).filter(RegistroMejora.id == rid).first()
    if not registro:
        raise HTTPException(status_code=404, detail="Registro no encontrado")
    for k, v in datos.model_dump(exclude_unset=True).items():
        setattr(registro, k, v)
    registro.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(registro)
    return registro

@app.post("/api/mejoras/{rid}/enviar", response_model=RegistroMejoraResponse)
def enviar_mejora(
    rid: int,
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    registro = db.query(RegistroMejora).filter(RegistroMejora.id == rid).first()
    if not registro:
        raise HTTPException(status_code=404, detail="Registro no encontrado")
    registro.estado = "enviado"
    registro.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(registro)
    return registro

@app.delete("/api/mejoras/{rid}", status_code=204)
def eliminar_mejora(
    rid: int,
    db: Session = Depends(get_db),
    admin: dict = Depends(verificar_admin)
):
    registro = db.query(RegistroMejora).filter(RegistroMejora.id == rid).first()
    if not registro:
        raise HTTPException(status_code=404, detail="Registro no encontrado")
    db.delete(registro)
    db.commit()

@app.get("/favicon.ico")
async def serve_favicon():
    """Servir el favicon"""
    favicon_path = os.path.join(FRONTEND_DIR, "favicon.ico")
    if os.path.exists(favicon_path):
        return FileResponse(favicon_path, media_type="image/x-icon")
    return None

@app.get("/")
async def serve_frontend():
    """Servir el frontend"""
    index_path = os.path.join(FRONTEND_DIR, "index.html")
    if os.path.exists(index_path):
        return FileResponse(index_path)
    return {"message": "Frontend no encontrado", "api_docs": "/docs"}

@app.get("/seguimiento")
async def serve_frontend_seguimiento():
    """URL directa a la vista de seguimiento — el JS se encarga de navegar al cargar."""
    index_path = os.path.join(FRONTEND_DIR, "index.html")
    if os.path.exists(index_path):
        return FileResponse(index_path)
    return {"message": "Frontend no encontrado"}


# ============================================
# PUNTO DE ENTRADA
# ============================================

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
